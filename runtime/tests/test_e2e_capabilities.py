"""Phase 9' — end-to-end capability smoke.

These tests run the real Python runtime (FastAPI app via TestClient) with
the LLM transport mocked at the `httpx.AsyncClient` boundary. Each test
verifies one wired-up capability:

  1. HumanInTheLoopMiddleware  — destructive tool triggers `run.waiting`
  2. SubAgentMiddleware        — `task` tool call surfaces `subagent.spawned`
  3. PromptCaching             — provider-specific cache markers stay outside
                                  the Runtime wire contract
  6. MemoryMiddleware          — AGENTS.md content lands in the outgoing
                                  system prompt
  7. TodoListMiddleware        — `write_todos` tool is in the agent toolset

Per capability: ~30 lines, one assertion per fact. They're "real path"
in everything except the LLM responses themselves.
"""

from __future__ import annotations

import asyncio
import json
import os
import tempfile
import threading
from pathlib import Path
from typing import Any

import httpx
import pytest
from docx import Document
from fastapi.testclient import TestClient

from shejane_runtime.config import reset_settings_for_tests
from shejane_runtime.server import create_app
from tests.helpers import run_command

# --- shared mock backend helpers ---


def _sse(events: list[tuple[str, dict[str, Any]]]) -> httpx.Response:
    body = "".join(
        f"event: {name}\ndata: {json.dumps(payload)}\n\n" for name, payload in events
    ).encode("utf-8")
    return httpx.Response(200, content=body, headers={"content-type": "text/event-stream"})


class RecordingHandler:
    """httpx.MockTransport callable that records every request body and
    returns a scripted SSE response based on call index.

    A list of canned responses is provided; each request pops the next.
    Falls back to a generic "done" stream if exhausted.
    """

    def __init__(self, scripts: list[list[tuple[str, dict[str, Any]]]]):
        self.scripts = list(scripts)
        self.requests: list[dict[str, Any]] = []
        self.completion_review_requests = 0
        self.title_generation_requests = 0

    @property
    def agent_requests(self) -> int:
        return len(self.requests) - self.completion_review_requests - self.title_generation_requests

    def __call__(self, request: httpx.Request) -> httpx.Response:
        body_bytes = request.read()
        try:
            body = json.loads(body_bytes)
            self.requests.append(body)
        except json.JSONDecodeError:
            body = {"raw": body_bytes.decode("utf-8", errors="replace")}
            self.requests.append(body)
        if "P9 final-answer reviewer" in str(body):
            self.completion_review_requests += 1
            return _sse(
                [
                    (
                        "llm.delta",
                        {
                            "content_delta": json.dumps(
                                {
                                    "decision": "allow",
                                    "reason": "The scripted final answer satisfies the test goal.",
                                }
                            )
                        },
                    ),
                    ("llm.done", {"request_id": "completion-review", "finish_reason": "stop"}),
                ]
            )
        if "conversation title generator" in str(body):
            self.title_generation_requests += 1
            return _sse(
                [
                    ("llm.delta", {"content_delta": "香港账户给美国 LLC 注资"}),
                    ("llm.done", {"request_id": "title", "finish_reason": "stop"}),
                ]
            )
        if self.scripts:
            return _sse(self.scripts.pop(0))
        return _sse([("llm.done", {"request_id": "x", "finish_reason": "stop"})])


def _patched_async_client(handler):
    class _Patched(httpx.AsyncClient):
        def __init__(self, **kw):
            transport = kw.pop("transport", None) or httpx.MockTransport(handler)
            super().__init__(
                transport=transport,
                **kw,
            )

    return _Patched


def _make_client(monkeypatch, handler) -> TestClient:
    tmp = Path(tempfile.mkdtemp(prefix="jdl-e2e-"))
    os.environ["SHEJANE_RUNTIME_TOKEN"] = "tok"
    monkeypatch.delenv("SHEJANE_RUNTIME_MCP_SERVERS", raising=False)
    monkeypatch.delenv("TAVILY_API_KEY", raising=False)
    monkeypatch.setattr("tests.streaming_model.httpx.AsyncClient", _patched_async_client(handler))
    settings = reset_settings_for_tests(
        SHEJANE_RUNTIME_HOST="127.0.0.1",
        SHEJANE_RUNTIME_PORT=17371,
        SHEJANE_RUNTIME_TOKEN="tok",
        data_dir=tmp,
    )
    app = create_app(settings)
    return TestClient(app)


def _parse_sse(body: str) -> list[tuple[str, dict[str, Any]]]:
    """Return list of (event_type, payload) — payload is auto-unwrapped
    from the AgentRunEvent envelope (Block 0). The `event:` framing
    line is dropped in favor of the envelope's `event_type`, so tests
    can keep asserting on payload contents (e.g. `data["goal"]`)."""
    events: list[tuple[str, dict[str, Any]]] = []
    name = ""
    buf: list[str] = []

    def flush() -> None:
        nonlocal name, buf
        if not (name or buf):
            return
        try:
            data = json.loads("\n".join(buf))
        except json.JSONDecodeError:
            data = {"raw": "\n".join(buf)}
        if isinstance(data, dict) and "event_type" in data and "payload" in data:
            events.append((str(data["event_type"]), data["payload"]))
        else:
            events.append((name, data))

    for raw in body.split("\n"):
        line = raw.rstrip("\r")
        if not line:
            flush()
            name = ""
            buf = []
            continue
        if line.startswith("event:"):
            name = line[6:].strip()
        elif line.startswith("data:"):
            buf.append(line[5:].strip())
    flush()
    return events


def _post_run_and_stream(
    client: TestClient,
    goal: str,
    *,
    workspace_path: str | None = None,
    **fields: Any,
) -> list[tuple[str, dict[str, Any]]]:
    body = run_command(goal, **fields)
    if workspace_path is not None:
        authorized = client.post(
            "/v1/workspaces",
            headers={"Authorization": "Bearer tok"},
            json={"path": workspace_path, "label": "test"},
        )
        assert authorized.status_code == 200, authorized.text
        body["workspace_path"] = workspace_path
    r = client.post(
        "/v1/runs",
        headers={"Authorization": "Bearer tok"},
        json=body,
    )
    assert r.status_code == 200, r.text
    run_id = r.json()["id"]
    with client.stream(
        "GET",
        f"/v1/runs/{run_id}/stream",
        headers={"Authorization": "Bearer tok"},
    ) as resp:
        body_text = resp.read().decode("utf-8")
    return _parse_sse(body_text)


def test_first_completed_turn_generates_runtime_owned_thread_title(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    handler = RecordingHandler(
        scripts=[
            [
                ("llm.delta", {"content_delta": "香港账户可以作为中转，但仍需合规申报。"}),
                ("llm.done", {"request_id": "answer", "finish_reason": "stop"}),
            ]
        ]
    )
    with _make_client(monkeypatch, handler) as client:
        events = _post_run_and_stream(
            client,
            "odi 听上去好复杂，还是香港比较好？帮我搜一下别人是怎么处理的",
            thread_id="thread_generated_title",
            assistant_message_id="msg_generated_title",
            thread_title="odi 听上去好复杂，还是香港比较好？",
        )
        snapshot = client.get(
            "/v1/threads/thread_generated_title",
            headers={"Authorization": "Bearer tok"},
        )

    assert any(name == "run.completed" for name, _payload in events)
    assert snapshot.status_code == 200
    assert snapshot.json()["thread"]["title"] == "香港账户给美国 LLC 注资"
    assert handler.title_generation_requests == 1


def test_run_reads_a_docx_snapshot_after_the_original_is_deleted(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    attachment = tmp_path / "contract.docx"
    document = Document()
    document.add_heading("Runtime-owned agreement", level=1)
    document.add_paragraph("Snapshot survives removal of the original file.")
    document.save(attachment)
    handler = RecordingHandler(
        scripts=[
            [
                (
                    "llm.tool_call",
                    {
                        "id": "call_read_docx_snapshot",
                        "name": "read_file",
                        "arguments": {"file_path": "/attachments/contract.docx"},
                    },
                ),
                ("llm.done", {"request_id": "read-docx", "finish_reason": "tool_calls"}),
            ],
            [
                (
                    "llm.tool_call",
                    {
                        "id": "call_office_read_docx_snapshot",
                        "name": "office.read",
                        "arguments": {"path": "/attachments/contract.docx"},
                    },
                ),
                ("llm.done", {"request_id": "office-docx", "finish_reason": "tool_calls"}),
            ],
            [
                ("llm.delta", {"content_delta": "Runtime-owned agreement"}),
                ("llm.done", {"request_id": "finish-docx", "finish_reason": "stop"}),
            ],
        ]
    )
    with _make_client(monkeypatch, handler) as client:
        body = run_command(
            "summarize the attachment",
            attachment_paths=[str(attachment)],
            required_capabilities=["agent.run", "agent.stream", "attachments"],
        )
        created = client.post(
            "/v1/runs",
            headers={"Authorization": "Bearer tok"},
            json=body,
        )
        assert created.status_code == 200, created.text
        attachment.unlink()

        with client.stream(
            "GET",
            f"/v1/runs/{created.json()['id']}/stream",
            headers={"Authorization": "Bearer tok"},
        ) as response:
            events = _parse_sse(response.read().decode("utf-8"))

    completed_tools = [payload for name, payload in events if name == "tool.completed"]
    assert {payload["name"] for payload in completed_tools} == {"read_file", "office.read"}
    assert "Snapshot survives removal" in json.dumps(handler.requests)
    office_result = next(payload for payload in completed_tools if payload["name"] == "office.read")
    assert '"ok": "true"' in office_result["content"]
    assert '"markdown"' in office_result["content"]
    completed = next(payload for name, payload in events if name == "run.completed")
    assert "Runtime-owned agreement" in completed["final_text"]


# ---- capability 1: HumanInTheLoop on destructive tool ----


def test_capability_1_humanintheloop_pauses_on_destructive_tool(monkeypatch) -> None:
    handler = RecordingHandler(
        scripts=[
            [
                # Mock LLM decides to call write_file (destructive).
                ("llm.delta", {"content_delta": "I'll write a file. "}),
                (
                    "llm.tool_call",
                    {
                        "id": "call_w1",
                        "name": "write_file",
                        "arguments": {"file_path": "spike.txt", "content": "hello"},
                    },
                ),
                ("llm.done", {"request_id": "r1", "finish_reason": "tool_calls"}),
            ]
        ]
    )
    with _make_client(monkeypatch, handler) as client:
        events = _post_run_and_stream(client, "please write a file")

    event_names = {e[0] for e in events}
    # HumanInTheLoop fires before write_file executes → graph pauses.
    assert "run.waiting" in event_names, (
        f"expected run.waiting (HumanInTheLoop interrupt). got: {events}"
    )
    # Block 4 contract: each pause must also surface a narrow
    # `permission.required` event carrying a `request_id` the client
    # can post back to /v1/permissions/{id}. Without this the UI
    # has no way to render an approval card.
    assert "permission.required" in event_names, (
        f"expected permission.required SSE event alongside run.waiting. got: {sorted(event_names)}"
    )
    perm_event = next(e for e in events if e[0] == "permission.required")
    perm_payload = perm_event[1]
    assert perm_payload["request_id"]
    assert perm_payload["tool"] == "write_file"
    assert perm_payload["allow_run_scope"] is True
    # `args` must round-trip — without them the approval card has no
    # context to show the user.
    assert perm_payload["arguments"]["file_path"] == "spike.txt"


def test_permission_mode_auto_executes_workspace_writes_without_prompt(
    monkeypatch,
    tmp_path,
) -> None:
    handler = RecordingHandler(
        scripts=[
            [
                (
                    "llm.tool_call",
                    {
                        "id": "call_auto_write",
                        "name": "write_file",
                        "arguments": {"file_path": "auto.txt", "content": "approved"},
                    },
                ),
                ("llm.done", {"request_id": "r1", "finish_reason": "tool_calls"}),
            ],
            [
                ("llm.delta", {"content_delta": "done"}),
                ("llm.done", {"request_id": "r2", "finish_reason": "stop"}),
            ],
        ]
    )
    with _make_client(monkeypatch, handler) as client:
        headers = {"Authorization": "Bearer tok"}
        authorized = client.post(
            "/v1/workspaces",
            headers=headers,
            json={"path": str(tmp_path), "label": "auto"},
        )
        assert authorized.status_code == 200, authorized.text
        command = run_command(
            "write without prompting",
            workspace_path=str(tmp_path),
            permission_mode="auto",
        )
        run = client.post("/v1/runs", headers=headers, json=command)
        assert run.status_code == 200, run.text
        with client.stream(
            "GET",
            f"/v1/runs/{run.json()['id']}/stream",
            headers=headers,
        ) as response:
            events = _parse_sse(response.read().decode("utf-8"))

    assert "run.completed" in {event[0] for event in events}
    assert "permission.required" not in {event[0] for event in events}
    assert (tmp_path / "auto.txt").read_text(encoding="utf-8") == "approved"


def test_workspace_write_without_workspace_fails_before_retry_loop(monkeypatch) -> None:
    handler = RecordingHandler(
        scripts=[
            [
                (
                    "llm.tool_call",
                    {
                        "id": "call_missing_workspace",
                        "name": "write_file",
                        "arguments": {"file_path": "result.html", "content": "done"},
                    },
                ),
                ("llm.done", {"request_id": "r1", "finish_reason": "tool_calls"}),
            ]
        ]
    )
    with _make_client(monkeypatch, handler) as client:
        events = _post_run_and_stream(
            client,
            "write a file",
            permission_mode="auto",
        )

    failed = next(payload for name, payload in events if name == "run.failed")
    assert failed["code"] == "workspace_required"
    assert failed["category"] == "workspace"
    assert failed["recovery_action"] == "workspace"
    assert handler.agent_requests == 1


def test_existing_file_conflict_is_structured_so_model_can_choose_a_new_name(
    monkeypatch,
    tmp_path,
) -> None:
    (tmp_path / "snake.html").write_text("existing", encoding="utf-8")
    (tmp_path / "snake-2.html").write_text("also existing", encoding="utf-8")

    def write(call_id: str) -> list[tuple[str, dict[str, Any]]]:
        return [
            (
                "llm.tool_call",
                {
                    "id": call_id,
                    "name": "write_file",
                    "arguments": {"file_path": "snake.html", "content": "replacement"},
                },
            ),
            ("llm.done", {"request_id": call_id, "finish_reason": "tool_calls"}),
        ]

    handler = RecordingHandler(
        scripts=[
            write("write-1"),
            [
                (
                    "llm.tool_call",
                    {
                        "id": "write-2",
                        "name": "write_file",
                        "arguments": {
                            "file_path": "snake-3.html",
                            "content": "replacement",
                        },
                    },
                ),
                ("llm.done", {"request_id": "write-2", "finish_reason": "tool_calls"}),
            ],
            [
                ("llm.delta", {"content_delta": "done"}),
                ("llm.done", {"request_id": "done", "finish_reason": "stop"}),
            ],
        ]
    )

    with _make_client(monkeypatch, handler) as client:
        events = _post_run_and_stream(
            client,
            "write snake.html",
            workspace_path=str(tmp_path),
            permission_mode="auto",
        )

    conflict = next(payload for name, payload in events if name == "tool.failed")
    conflict_content = json.loads(conflict["content"])
    assert conflict["error_code"] == "file_exists"
    assert conflict["recoverable"] is True
    assert conflict["retryable"] is False
    assert conflict_content["allowed_actions"] == [
        "choose_new_path",
        "read_then_edit",
        "ask_user",
    ]
    assert conflict_content["suggested_path"] == "snake-3.html"
    assert "run.completed" in {name for name, _ in events}
    assert sum(name == "tool.failed" for name, _ in events) == 1
    assert sum(name == "tool.completed" for name, _ in events) == 1
    assert (tmp_path / "snake.html").read_text(encoding="utf-8") == "existing"
    assert (tmp_path / "snake-2.html").read_text(encoding="utf-8") == "also existing"
    assert (tmp_path / "snake-3.html").read_text(encoding="utf-8") == "replacement"
    assert handler.agent_requests == 3


def test_read_file_without_pagination_reads_a_normal_text_file_in_one_call(
    monkeypatch,
    tmp_path,
) -> None:
    (tmp_path / "notes.txt").write_text(
        "".join(f"line {index}\n" for index in range(1, 151)),
        encoding="utf-8",
    )
    handler = RecordingHandler(
        scripts=[
            [
                (
                    "llm.tool_call",
                    {
                        "id": "read-1",
                        "name": "read_file",
                        "arguments": {"file_path": "notes.txt"},
                    },
                ),
                ("llm.done", {"request_id": "read-1", "finish_reason": "tool_calls"}),
            ],
            [
                ("llm.delta", {"content_delta": "done"}),
                ("llm.done", {"request_id": "done", "finish_reason": "stop"}),
            ],
        ]
    )

    with _make_client(monkeypatch, handler) as client:
        events = _post_run_and_stream(
            client,
            "read notes.txt",
            workspace_path=str(tmp_path),
            permission_mode="auto",
        )

    result = next(payload for name, payload in events if name == "tool.completed")
    assert "150\tline 150" in result["content"]
    assert handler.agent_requests == 2


def test_repeated_same_path_conflict_asks_the_user_before_more_writes(
    monkeypatch,
    tmp_path,
) -> None:
    (tmp_path / "snake.html").write_text("existing", encoding="utf-8")

    def write(call_id: str) -> list[tuple[str, dict[str, Any]]]:
        return [
            (
                "llm.tool_call",
                {
                    "id": call_id,
                    "name": "write_file",
                    "arguments": {"file_path": "snake.html", "content": "replacement"},
                },
            ),
            ("llm.done", {"request_id": call_id, "finish_reason": "tool_calls"}),
        ]

    handler = RecordingHandler(
        scripts=[
            write("write-1"),
            write("write-2"),
        ]
    )

    with _make_client(monkeypatch, handler) as client:
        events = _post_run_and_stream(
            client,
            "write snake.html",
            workspace_path=str(tmp_path),
            permission_mode="auto",
        )

    question = next(payload for name, payload in events if name == "question.asked")
    assert question["questions"][0]["question"] == "snake.html 已存在，如何处理？"
    assert [option["label"] for option in question["questions"][0]["options"]] == [
        "自动换名",
        "覆盖原文件",
        "取消写入",
    ]
    assert "run.waiting" in {name for name, _ in events}
    assert not any(name == "run.failed" for name, _ in events)
    assert sum(name == "tool.failed" for name, _ in events) == 1
    assert (tmp_path / "snake.html").read_text(encoding="utf-8") == "existing"
    assert handler.agent_requests == 2


def test_interleaved_file_conflicts_still_ask_before_retrying_the_same_path(
    monkeypatch,
    tmp_path,
) -> None:
    (tmp_path / "snake.html").write_text("existing", encoding="utf-8")
    (tmp_path / "贪吃蛇.html").write_text("existing", encoding="utf-8")

    def write(call_id: str, file_path: str) -> list[tuple[str, dict[str, Any]]]:
        return [
            (
                "llm.tool_call",
                {
                    "id": call_id,
                    "name": "write_file",
                    "arguments": {"file_path": file_path, "content": "replacement"},
                },
            ),
            ("llm.done", {"request_id": call_id, "finish_reason": "tool_calls"}),
        ]

    handler = RecordingHandler(
        scripts=[
            write("write-1", "snake.html"),
            write("write-2", "贪吃蛇.html"),
            write("write-3", "snake.html"),
        ]
    )

    with _make_client(monkeypatch, handler) as client:
        events = _post_run_and_stream(
            client,
            "write snake.html",
            workspace_path=str(tmp_path),
            permission_mode="auto",
        )

    assert any(name == "question.asked" for name, _ in events)
    assert sum(name == "tool.failed" for name, _ in events) == 2
    assert handler.agent_requests == 3


@pytest.mark.parametrize(
    ("answer", "expected_original", "expected_renamed"),
    [
        ("自动换名", "existing", "replacement"),
        ("覆盖原文件", "replacement", None),
        ("取消写入", "existing", None),
    ],
)
def test_file_conflict_answer_resumes_the_paused_write(
    monkeypatch,
    tmp_path,
    answer: str,
    expected_original: str,
    expected_renamed: str | None,
) -> None:
    (tmp_path / "snake.html").write_text("existing", encoding="utf-8")

    def write(call_id: str) -> list[tuple[str, dict[str, Any]]]:
        return [
            (
                "llm.tool_call",
                {
                    "id": call_id,
                    "name": "write_file",
                    "arguments": {"file_path": "snake.html", "content": "replacement"},
                },
            ),
            ("llm.done", {"request_id": call_id, "finish_reason": "tool_calls"}),
        ]

    handler = RecordingHandler(
        scripts=[
            write("write-1"),
            write("write-2"),
            [
                ("llm.delta", {"content_delta": "done"}),
                ("llm.done", {"request_id": "done", "finish_reason": "stop"}),
            ],
        ]
    )

    with _make_client(monkeypatch, handler) as client:
        headers = {"Authorization": "Bearer tok"}
        authorized = client.post(
            "/v1/workspaces",
            headers=headers,
            json={"path": str(tmp_path), "label": "test"},
        )
        assert authorized.status_code == 200, authorized.text
        run = client.post(
            "/v1/runs",
            headers=headers,
            json=run_command(
                "write snake.html",
                workspace_path=str(tmp_path),
                permission_mode="auto",
            ),
        )
        assert run.status_code == 200, run.text
        run_id = run.json()["id"]

        with client.stream("GET", f"/v1/runs/{run_id}/stream", headers=headers) as resp:
            first_events = _parse_sse(resp.read().decode("utf-8"))
        question = next(payload for name, payload in first_events if name == "question.asked")
        question_id = question["request_id"]
        resumed = client.post(
            f"/v1/questions/{question_id}",
            headers=headers,
            json={"answers": {question_id: [answer]}},
        )
        assert resumed.status_code == 200, resumed.text

        with client.stream("GET", f"/v1/runs/{run_id}/stream", headers=headers) as resp:
            resumed_events = _parse_sse(resp.read().decode("utf-8"))

    assert "run.completed" in {name for name, _ in resumed_events}
    assert (tmp_path / "snake.html").read_text(encoding="utf-8") == expected_original
    renamed = tmp_path / "snake-2.html"
    if expected_renamed is None:
        assert not renamed.exists()
    else:
        assert renamed.read_text(encoding="utf-8") == expected_renamed


def test_auto_rename_redirects_later_file_tools_without_asking_again(
    monkeypatch,
    tmp_path,
) -> None:
    (tmp_path / "snake.html").write_text("existing", encoding="utf-8")

    def tool_call(call_id: str, name: str, arguments: dict[str, Any]):
        return [
            (
                "llm.tool_call",
                {"id": call_id, "name": name, "arguments": arguments},
            ),
            ("llm.done", {"request_id": call_id, "finish_reason": "tool_calls"}),
        ]

    original = str(tmp_path / "snake.html")
    handler = RecordingHandler(
        scripts=[
            tool_call("write-1", "write_file", {"file_path": original, "content": "draft"}),
            tool_call("write-2", "write_file", {"file_path": original, "content": "draft"}),
            tool_call("read-1", "read_file", {"file_path": original}),
            tool_call(
                "edit-1",
                "edit_file",
                {"file_path": original, "old_string": "draft", "new_string": "refined"},
            ),
            tool_call("write-3", "write_file", {"file_path": original, "content": "ignored"}),
            [
                ("llm.delta", {"content_delta": "done"}),
                ("llm.done", {"request_id": "done", "finish_reason": "stop"}),
            ],
        ]
    )

    with _make_client(monkeypatch, handler) as client:
        headers = {"Authorization": "Bearer tok"}
        authorized = client.post(
            "/v1/workspaces",
            headers=headers,
            json={"path": str(tmp_path), "label": "test"},
        )
        assert authorized.status_code == 200, authorized.text
        run = client.post(
            "/v1/runs",
            headers=headers,
            json=run_command(
                "write snake.html",
                workspace_path=str(tmp_path),
                permission_mode="auto",
            ),
        )
        run_id = run.json()["id"]

        with client.stream("GET", f"/v1/runs/{run_id}/stream", headers=headers) as resp:
            first_events = _parse_sse(resp.read().decode("utf-8"))
        question = next(payload for name, payload in first_events if name == "question.asked")
        question_id = question["request_id"]
        resumed = client.post(
            f"/v1/questions/{question_id}",
            headers=headers,
            json={"answers": {question_id: ["自动换名"]}},
        )
        assert resumed.status_code == 200, resumed.text

        with client.stream("GET", f"/v1/runs/{run_id}/stream", headers=headers) as resp:
            resumed_events = _parse_sse(resp.read().decode("utf-8"))

    all_events = first_events + resumed_events
    assert "run.completed" in {name for name, _ in resumed_events}
    assert {payload["request_id"] for name, payload in all_events if name == "question.asked"} == {
        question_id
    }
    assert (tmp_path / "snake.html").read_text(encoding="utf-8") == "existing"
    assert (tmp_path / "snake-2.html").read_text(encoding="utf-8") == "refined"


def test_permission_mode_auto_allows_sandboxed_command_without_model_review(
    monkeypatch,
    tmp_path,
) -> None:
    class ApprovalAwareHandler(RecordingHandler):
        def __call__(self, request: httpx.Request) -> httpx.Response:
            body = json.loads(request.read())
            self.requests.append(body)
            if len(self.requests) == 1:
                return _sse(
                    [
                        (
                            "llm.tool_call",
                            {
                                "id": "call_auto_execute",
                                "name": "execute",
                                "arguments": {"command": "cat reviewed.txt"},
                            },
                        ),
                        ("llm.done", {"request_id": "r1", "finish_reason": "tool_calls"}),
                    ]
                )
            return _sse(
                [
                    ("llm.delta", {"content_delta": "done"}),
                    ("llm.done", {"request_id": "r2", "finish_reason": "stop"}),
                ]
            )

    handler = ApprovalAwareHandler(scripts=[])
    with _make_client(monkeypatch, handler) as client:
        (tmp_path / "reviewed.txt").write_text("reviewed", encoding="utf-8")
        headers = {"Authorization": "Bearer tok"}
        authorized = client.post(
            "/v1/workspaces",
            headers=headers,
            json={"path": str(tmp_path), "label": "auto-review"},
        )
        assert authorized.status_code == 200, authorized.text
        run = client.post(
            "/v1/runs",
            headers=headers,
            json=run_command(
                "read reviewed.txt in the workspace",
                workspace_path=str(tmp_path),
                permission_mode="auto",
            ),
        )
        assert run.status_code == 200, run.text
        run_id = run.json()["id"]
        with client.stream(
            "GET",
            f"/v1/runs/{run_id}/stream",
            headers=headers,
        ) as response:
            events = _parse_sse(response.read().decode("utf-8"))
        receipts = client.portal.call(client.app.state.store.list_tool_receipts_for_run, run_id)

    approval = next(payload for name, payload in events if name == "permission.auto_approved")
    assert approval["source"] == "rule"
    assert "permission.required" not in {name for name, _payload in events}
    assert (tmp_path / "reviewed.txt").read_text(encoding="utf-8") == "reviewed"
    assert receipts[0]["review_source"] == "rule"


def test_permission_mode_auto_still_prompts_for_sensitive_and_external_tools(monkeypatch) -> None:
    handler = RecordingHandler(
        scripts=[
            [
                (
                    "llm.tool_call",
                    {"id": "call_auto_clipboard", "name": "clipboard.read", "arguments": {}},
                ),
                (
                    "llm.tool_call",
                    {
                        "id": "call_auto_external",
                        "name": "clipboard.write",
                        "arguments": {"text": "do not write without approval"},
                    },
                ),
                ("llm.done", {"request_id": "r1", "finish_reason": "tool_calls"}),
            ]
        ]
    )
    with _make_client(monkeypatch, handler) as client:
        events = _post_run_and_stream(
            client,
            "read clipboard",
            permission_mode="auto",
        )

    required = [event for event in events if event[0] == "permission.required"]
    assert {event[1]["tool"] for event in required} == {"clipboard.read", "clipboard.write"}
    external = next(event[1] for event in required if event[1]["tool"] == "clipboard.write")
    assert external["review_source"] == "fallback"
    assert "fallback policy" in external["review_reason"]


def test_permission_mode_full_access_executes_clipboard_read_without_prompt(monkeypatch) -> None:
    handler = RecordingHandler(
        scripts=[
            [
                (
                    "llm.tool_call",
                    {"id": "call_full_clipboard", "name": "clipboard.read", "arguments": {}},
                ),
                ("llm.done", {"request_id": "r1", "finish_reason": "tool_calls"}),
            ],
            [
                ("llm.delta", {"content_delta": "done"}),
                ("llm.done", {"request_id": "r2", "finish_reason": "stop"}),
            ],
        ]
    )
    with _make_client(monkeypatch, handler) as client:
        events = _post_run_and_stream(
            client,
            "read clipboard",
            permission_mode="full_access",
        )

    names = {event[0] for event in events}
    assert "run.completed" in names
    assert "permission.required" not in names


def test_capability_1c_permission_resolved_event_clears_card(monkeypatch) -> None:
    """POST /permissions/:id must emit `permission.resolved` onto the
    SSE stream so the client's `hasPendingPermission` set (App.tsx:1339)
    drops the request_id and the approval card disappears. Without this
    the card stays visible after the user clicks approve even though
    the run has already moved on.
    """
    handler = RecordingHandler(
        scripts=[
            [
                (
                    "llm.tool_call",
                    {
                        "id": "call_w1",
                        "name": "write_file",
                        "arguments": {"file_path": "x.txt", "content": "hi"},
                    },
                ),
                ("llm.done", {"request_id": "r1", "finish_reason": "tool_calls"}),
            ],
            [
                ("llm.delta", {"content_delta": "ok"}),
                ("llm.done", {"request_id": "r2", "finish_reason": "stop"}),
            ],
        ]
    )
    with _make_client(monkeypatch, handler) as client:
        headers = {"Authorization": "Bearer tok"}
        run = client.post("/v1/runs", headers=headers, json=run_command("write")).json()
        with client.stream("GET", f"/v1/runs/{run['id']}/stream", headers=headers) as resp:
            resp.read()
        perm_id = next(
            e for e in _parse_sse_persisted(client, run["id"]) if e[0] == "permission.required"
        )[1]["request_id"]
        client.post(
            f"/v1/permissions/{perm_id}",
            headers=headers,
            json={"decision": "approve", "scope": "once"},
        )
        with client.stream("GET", f"/v1/runs/{run['id']}/stream", headers=headers) as resp:
            body = resp.read().decode("utf-8")
    events = _parse_sse(body)
    resolved = [e for e in events if e[0] == "permission.resolved"]
    assert resolved, "expected permission.resolved on the post-resume stream"
    assert resolved[0][1]["request_id"] == perm_id
    assert resolved[0][1]["decision"] == "approve"


def _parse_sse_persisted(client: TestClient, run_id: str) -> list[tuple[str, dict[str, Any]]]:
    """Fetch all persisted events for a run via /diagnostics — used by
    tests that need the pre-pause event list without re-streaming."""
    diag = client.get(
        f"/v1/runs/{run_id}/diagnostics",
        headers={"Authorization": "Bearer tok"},
    ).json()
    return [(e["event_type"], e["payload"]) for e in diag["events"]]


def test_capability_1d_scope_run_stops_asking_for_new_arguments(monkeypatch, tmp_path) -> None:
    """ "Don't ask again" covers later ordinary calls to the same tool."""
    handler = RecordingHandler(
        scripts=[
            # Turn 1: ask to write file A (paused by HITL)
            [
                (
                    "llm.tool_call",
                    {
                        "id": "call_a",
                        "name": "write_file",
                        "arguments": {"file_path": "a.txt", "content": "A"},
                    },
                ),
                ("llm.done", {"request_id": "r1", "finish_reason": "tool_calls"}),
            ],
            # Turn 2 (after "don't ask again"): a different path executes
            # without another permission pause.
            [
                (
                    "llm.tool_call",
                    {
                        "id": "call_b",
                        "name": "write_file",
                        "arguments": {"file_path": "b.txt", "content": "B"},
                    },
                ),
                ("llm.done", {"request_id": "r2", "finish_reason": "tool_calls"}),
            ],
            # Turn 3: final answer
            [
                ("llm.delta", {"content_delta": "done"}),
                ("llm.done", {"request_id": "r3", "finish_reason": "stop"}),
            ],
        ]
    )
    with _make_client(monkeypatch, handler) as client:
        headers = {"Authorization": "Bearer tok"}
        authorized = client.post(
            "/v1/workspaces",
            headers=headers,
            json={"path": str(tmp_path), "label": "permissions"},
        )
        assert authorized.status_code == 200, authorized.text
        run = client.post(
            "/v1/runs",
            headers=headers,
            json=run_command("write two files", workspace_path=str(tmp_path)),
        ).json()
        with client.stream("GET", f"/v1/runs/{run['id']}/stream", headers=headers) as resp:
            resp.read()
        # Approve this ordinary tool for the rest of the run.
        perm_id = next(
            e for e in _parse_sse_persisted(client, run["id"]) if e[0] == "permission.required"
        )[1]["request_id"]
        first_approval = client.post(
            f"/v1/permissions/{perm_id}",
            headers=headers,
            json={"decision": "approve", "scope": "run"},
        )
        assert first_approval.status_code == 200, first_approval.text
        with client.stream("GET", f"/v1/runs/{run['id']}/stream", headers=headers) as resp:
            body = resp.read().decode("utf-8")
        events = _parse_sse(body)
        # A new subscriber replays the durable log from its own cursor. The
        # second call must execute without creating a new review candidate.
        second_required = [
            e
            for e in events
            if e[0] == "permission.required" and e[1].get("tool_call_id") == "call_b"
        ]
        assert second_required == []

    assert "run.completed" in [event[0] for event in events]
    assert (tmp_path / "a.txt").read_text() == "A"
    assert (tmp_path / "b.txt").read_text() == "B"


def test_capability_1b_permission_approve_resumes_the_run(monkeypatch, tmp_path) -> None:
    """Full pause → POST /permissions/:id → resume cycle.

    Regression for the `decisions = interrupt(hitl_request)["decisions"]`
    KeyError: HumanInTheLoopMiddleware (langchain.agents.middleware) only
    accepts `Command(resume={"decisions": [{"type": "approve"|"reject"|...}]})`.
    Our `POST /v1/permissions/{id}` must translate the client's
    `{decision: 'approve'|'deny'}` body into that shape — if it passes
    the raw client payload through, the middleware crashes mid-resume
    and the run dies with no AI response."""
    handler = RecordingHandler(
        scripts=[
            # Turn 1: model asks to write a file (paused by HITL)
            [
                (
                    "llm.tool_call",
                    {
                        "id": "call_w1",
                        "name": "write_file",
                        "arguments": {"file_path": "ok.txt", "content": "x"},
                    },
                ),
                ("llm.done", {"request_id": "r1", "finish_reason": "tool_calls"}),
            ],
            # Turn 2 (after approve + tool exec): final answer
            [
                ("llm.delta", {"content_delta": "Wrote it."}),
                ("llm.done", {"request_id": "r2", "finish_reason": "stop"}),
            ],
        ]
    )
    with _make_client(monkeypatch, handler) as client:
        headers = {"Authorization": "Bearer tok"}
        authorized = client.post(
            "/v1/workspaces",
            headers=headers,
            json={"path": str(tmp_path), "label": "permissions"},
        )
        assert authorized.status_code == 200, authorized.text
        r = client.post(
            "/v1/runs",
            headers=headers,
            json=run_command("write a file", workspace_path=str(tmp_path)),
        )
        run_id = r.json()["id"]
        with client.stream("GET", f"/v1/runs/{run_id}/stream", headers=headers) as resp:
            body1 = resp.read().decode("utf-8")
        events1 = _parse_sse(body1)
        perm = next(e for e in events1 if e[0] == "permission.required")
        permission_id = perm[1]["request_id"]

        # Client-shape POST — must NOT include `run_id`, must use the
        # `decision/scope` keys per client.ts:resolveLocalPermission.
        approve = client.post(
            f"/v1/permissions/{permission_id}",
            headers=headers,
            json={"decision": "approve", "scope": "once"},
        )
        assert approve.status_code == 200, approve.text
        assert approve.json()["resolved"] is True

        # Drain post-resume stream — should reach run.completed (not
        # run.failed with KeyError: 'decisions').
        with client.stream("GET", f"/v1/runs/{run_id}/stream", headers=headers) as resp:
            body2 = resp.read().decode("utf-8")

    events2 = _parse_sse(body2)
    names2 = {e[0] for e in events2}
    assert "run.completed" in names2, f"expected run.completed after approve. got: {sorted(names2)}"
    assert "run.failed" not in names2, "approve resume should not crash the graph"


def test_capability_1e_denied_tool_is_not_executed_and_has_rejected_receipt(
    monkeypatch,
) -> None:
    handler = RecordingHandler(
        scripts=[
            [
                (
                    "llm.tool_call",
                    {
                        "id": "call_denied",
                        "name": "write_file",
                        "arguments": {"file_path": "denied.txt", "content": "no"},
                    },
                ),
                ("llm.done", {"request_id": "r1", "finish_reason": "tool_calls"}),
            ],
            [
                ("llm.delta", {"content_delta": "The write was denied."}),
                ("llm.done", {"request_id": "r2", "finish_reason": "stop"}),
            ],
        ]
    )
    with _make_client(monkeypatch, handler) as client:
        headers = {"Authorization": "Bearer tok"}
        run = client.post("/v1/runs", headers=headers, json=run_command("do not write")).json()
        with client.stream("GET", f"/v1/runs/{run['id']}/stream", headers=headers) as response:
            response.read()
        permission = next(
            event
            for event in _parse_sse_persisted(client, run["id"])
            if event[0] == "permission.required"
        )[1]
        denied = client.post(
            f"/v1/permissions/{permission['request_id']}",
            headers=headers,
            json={"decision": "deny", "scope": "once"},
        )
        assert denied.status_code == 200, denied.text
        with client.stream("GET", f"/v1/runs/{run['id']}/stream", headers=headers) as response:
            body = response.read().decode("utf-8")
        diagnostics = client.get(f"/v1/runs/{run['id']}/diagnostics", headers=headers).json()

    assert "run.completed" in [event[0] for event in _parse_sse(body)], body
    receipt = next(
        item for item in diagnostics["tool_receipts"] if item["tool_call_id"] == "call_denied"
    )
    assert receipt["status"] == "rejected"
    assert receipt["attempt_count"] == 0


def test_capability_1f_review_pauses_the_entire_mixed_tool_batch(
    monkeypatch,
    tmp_path,
) -> None:
    (tmp_path / "source.txt").write_text("source", encoding="utf-8")
    handler = RecordingHandler(
        scripts=[
            [
                (
                    "llm.tool_call",
                    {
                        "id": "call_read",
                        "name": "read_file",
                        "arguments": {"file_path": "source.txt"},
                    },
                ),
                (
                    "llm.tool_call",
                    {
                        "id": "call_write",
                        "name": "write_file",
                        "arguments": {"file_path": "target.txt", "content": "target"},
                    },
                ),
                ("llm.done", {"request_id": "r1", "finish_reason": "tool_calls"}),
            ],
            [
                ("llm.delta", {"content_delta": "Both calls finished."}),
                ("llm.done", {"request_id": "r2", "finish_reason": "stop"}),
            ],
        ]
    )
    with _make_client(monkeypatch, handler) as client:
        headers = {"Authorization": "Bearer tok"}
        workspace = client.post(
            "/v1/workspaces",
            headers=headers,
            json={"path": str(tmp_path), "label": "batch"},
        )
        assert workspace.status_code == 200, workspace.text
        command = run_command("read then write")
        command["workspace_path"] = str(tmp_path)
        run = client.post("/v1/runs", headers=headers, json=command).json()
        with client.stream("GET", f"/v1/runs/{run['id']}/stream", headers=headers) as response:
            response.read()

        paused = client.get(f"/v1/runs/{run['id']}/diagnostics", headers=headers).json()
        # Review now prepares the full batch durably, but no sibling may start
        # before the consequential call has resolved.
        assert len(paused["tool_receipts"]) == 2
        assert {receipt["status"] for receipt in paused["tool_receipts"]} == {"prepared"}
        assert {receipt["attempt_count"] for receipt in paused["tool_receipts"]} == {0}
        permission = next(
            event for event in paused["events"] if event["event_type"] == "permission.required"
        )["payload"]
        approved = client.post(
            f"/v1/permissions/{permission['request_id']}",
            headers=headers,
            json={"decision": "approve", "scope": "once"},
        )
        assert approved.status_code == 200, approved.text
        with client.stream("GET", f"/v1/runs/{run['id']}/stream", headers=headers) as response:
            response.read()
        completed = client.get(f"/v1/runs/{run['id']}/diagnostics", headers=headers).json()

    assert {receipt["tool_call_id"] for receipt in completed["tool_receipts"]} == {
        "call_read",
        "call_write",
    }
    assert (tmp_path / "target.txt").read_text(encoding="utf-8") == "target"


def test_capability_1g_invalid_tool_arguments_fail_before_review(monkeypatch) -> None:
    handler = RecordingHandler(
        scripts=[
            [
                (
                    "llm.tool_call",
                    {
                        "id": "call_invalid",
                        "name": "write_file",
                        # `content` is required; this obsolete key must be
                        # rejected before asking the user or entering a tool.
                        "arguments": {"file_path": "bad.txt", "text": "bad"},
                    },
                ),
                ("llm.done", {"request_id": "r1", "finish_reason": "tool_calls"}),
            ],
            [
                ("llm.delta", {"content_delta": "The call was invalid."}),
                ("llm.done", {"request_id": "r2", "finish_reason": "stop"}),
            ],
        ]
    )
    with _make_client(monkeypatch, handler) as client:
        headers = {"Authorization": "Bearer tok"}
        run = client.post("/v1/runs", headers=headers, json=run_command("invalid write")).json()
        with client.stream("GET", f"/v1/runs/{run['id']}/stream", headers=headers) as response:
            body = response.read().decode("utf-8")
        diagnostics = client.get(f"/v1/runs/{run['id']}/diagnostics", headers=headers).json()

    assert "run.completed" in [event[0] for event in _parse_sse(body)]
    assert not any(event["event_type"] == "permission.required" for event in diagnostics["events"])
    receipt = next(
        item for item in diagnostics["tool_receipts"] if item["tool_call_id"] == "call_invalid"
    )
    assert receipt["status"] == "failed"
    assert receipt["attempt_count"] == 0
    assert receipt["error_type"] == "ToolInputValidationError"


def test_capability_1h_edited_arguments_are_revalidated_and_executed(
    monkeypatch,
    tmp_path,
) -> None:
    handler = RecordingHandler(
        scripts=[
            [
                (
                    "llm.tool_call",
                    {
                        "id": "call_edit_write",
                        "name": "write_file",
                        "arguments": {"file_path": "edited.txt", "content": "original"},
                    },
                ),
                ("llm.done", {"request_id": "r1", "finish_reason": "tool_calls"}),
            ],
            [
                ("llm.delta", {"content_delta": "Edited write finished."}),
                ("llm.done", {"request_id": "r2", "finish_reason": "stop"}),
            ],
        ]
    )
    with _make_client(monkeypatch, handler) as client:
        headers = {"Authorization": "Bearer tok"}
        client.post(
            "/v1/workspaces",
            headers=headers,
            json={"path": str(tmp_path), "label": "edit"},
        )
        command = run_command("write edited content")
        command["workspace_path"] = str(tmp_path)
        run = client.post("/v1/runs", headers=headers, json=command).json()
        with client.stream("GET", f"/v1/runs/{run['id']}/stream", headers=headers) as response:
            first_body = response.read().decode("utf-8")
        permission = next(
            event[1] for event in _parse_sse(first_body) if event[0] == "permission.required"
        )
        edited_action = {
            "name": "write_file",
            "args": {"file_path": "edited.txt", "content": "edited"},
        }
        edited = client.post(
            f"/v1/permissions/{permission['request_id']}",
            headers=headers,
            json={
                "decision": "edit",
                "scope": "once",
                "edited_action": edited_action,
            },
        )
        assert edited.status_code == 200, edited.text
        with client.stream("GET", f"/v1/runs/{run['id']}/stream", headers=headers) as response:
            second_body = response.read().decode("utf-8")

    assert "run.completed" in [event[0] for event in _parse_sse(second_body)]
    assert (tmp_path / "edited.txt").read_text(encoding="utf-8") == "edited"


# ---- capability 2: SubAgent dispatch ----


def test_capability_2_subagent_task_surfaces_spawned_event(monkeypatch) -> None:
    handler = RecordingHandler(
        scripts=[
            [
                ("llm.delta", {"content_delta": "Let me research that. "}),
                (
                    "llm.tool_call",
                    {
                        "id": "call_t1",
                        "name": "task",
                        "arguments": {
                            "subagent_type": "researcher",
                            "description": "find the latest LangGraph release notes",
                        },
                    },
                ),
                ("llm.done", {"request_id": "r2", "finish_reason": "tool_calls"}),
            ],
            # Researcher subagent's LLM turn — return a short finding then done.
            [
                ("llm.delta", {"content_delta": "Found notes."}),
                ("llm.done", {"request_id": "r3", "finish_reason": "stop"}),
            ],
        ]
    )
    with _make_client(monkeypatch, handler) as client:
        events = _post_run_and_stream(client, "research the latest LangGraph notes")

    spawn = next((payload for name, payload in events if name == "subagent.spawned"), None)
    assert spawn is not None, f"expected durable task receipt spawn. got: {events}"
    assert spawn["operation_id"].startswith("toolop_")
    assert spawn["tool_call_id"] == "call_t1"
    assert spawn["status"] == "queued"


def test_subagent_transient_model_failure_is_retried_and_contained(monkeypatch) -> None:
    from shejane_runtime.llm import ledger

    original_retry_decision = ledger.build_retry_decision

    def immediate_retry(*args, **kwargs):
        return {**original_retry_decision(*args, **kwargs), "delay_s": 0.0}

    monkeypatch.setattr(ledger, "build_retry_decision", immediate_retry)
    unavailable = [
        (
            "llm.error",
            {
                "request_id": "provider-unavailable",
                "message": "Service temporarily unavailable",
                "code": "service_unavailable",
                "recoverable": True,
                "retryable": True,
            },
        )
    ]
    handler = RecordingHandler(
        scripts=[
            [
                (
                    "llm.tool_call",
                    {
                        "id": "call_retrying_task",
                        "name": "task",
                        "arguments": {
                            "subagent_type": "researcher",
                            "description": "Check one source",
                        },
                    },
                ),
                ("llm.done", {"request_id": "parent-1", "finish_reason": "tool_calls"}),
            ],
            unavailable,
            unavailable,
            unavailable,
            [
                ("llm.delta", {"content_delta": "I could not reach the source, so I stopped."}),
                ("llm.done", {"request_id": "parent-2", "finish_reason": "stop"}),
            ],
        ]
    )

    with _make_client(monkeypatch, handler) as client:
        headers = {"Authorization": "Bearer tok"}
        run = client.post(
            "/v1/runs",
            headers=headers,
            json=run_command("Use the task tool to delegate this question to the researcher"),
        ).json()
        with client.stream(
            "GET",
            f"/v1/runs/{run['id']}/stream",
            headers=headers,
        ) as response:
            events = _parse_sse(response.read().decode("utf-8"))
        diagnostics = client.get(
            f"/v1/runs/{run['id']}/diagnostics",
            headers=headers,
        ).json()

    assert "run.completed" in {name for name, _payload in events}
    assert "ExecutionSettlementError" not in json.dumps(events)
    task_receipt = next(
        receipt for receipt in diagnostics["tool_receipts"] if receipt["tool_name"] == "task"
    )
    assert task_receipt["status"] == "failed"
    retries = [
        call
        for call in diagnostics["model_calls"]
        if call["provider_request_id"] == "provider-unavailable"
    ]
    assert [call["retry_attempt"] for call in retries] == [0, 1, 2]
    assert len({call["logical_call_id"] for call in retries}) == 1


def test_failed_subagent_can_be_replaced_after_two_prior_task_attempts(monkeypatch) -> None:
    from shejane_runtime.llm import ledger

    original_retry_decision = ledger.build_retry_decision

    def immediate_retry(*args, **kwargs):
        return {**original_retry_decision(*args, **kwargs), "delay_s": 0.0}

    monkeypatch.setattr(ledger, "build_retry_decision", immediate_retry)
    unavailable = [
        (
            "llm.error",
            {
                "request_id": "replacement-provider-unavailable",
                "message": "Service temporarily unavailable",
                "code": "service_unavailable",
                "recoverable": True,
                "retryable": True,
            },
        )
    ]
    handler = RecordingHandler(
        scripts=[
            [
                (
                    "llm.tool_call",
                    {
                        "id": "call_first_task",
                        "name": "task",
                        "arguments": {
                            "subagent_type": "researcher",
                            "description": "Check the first source",
                        },
                    },
                ),
                ("llm.done", {"request_id": "parent-1", "finish_reason": "tool_calls"}),
            ],
            [
                ("llm.delta", {"content_delta": "First source checked."}),
                ("llm.done", {"request_id": "child-1", "finish_reason": "stop"}),
            ],
            [
                (
                    "llm.tool_call",
                    {
                        "id": "call_failed_task",
                        "name": "task",
                        "arguments": {
                            "subagent_type": "researcher",
                            "description": "Check the second source",
                        },
                    },
                ),
                ("llm.done", {"request_id": "parent-2", "finish_reason": "tool_calls"}),
            ],
            unavailable,
            unavailable,
            unavailable,
            [
                (
                    "llm.tool_call",
                    {
                        "id": "call_replacement_task",
                        "name": "task",
                        "arguments": {
                            "subagent_type": "researcher",
                            "description": "Replace the failed second-source check",
                        },
                    },
                ),
                ("llm.done", {"request_id": "parent-3", "finish_reason": "tool_calls"}),
            ],
            [
                ("llm.delta", {"content_delta": "Replacement source checked."}),
                ("llm.done", {"request_id": "child-3", "finish_reason": "stop"}),
            ],
            [
                ("llm.delta", {"content_delta": "Replacement completed the research."}),
                ("llm.done", {"request_id": "parent-final", "finish_reason": "stop"}),
            ],
        ]
    )

    with _make_client(monkeypatch, handler) as client:
        headers = {"Authorization": "Bearer tok"}
        run = client.post(
            "/v1/runs",
            headers=headers,
            json=run_command(
                "Research independent sources and replace a transiently failed subagent"
            ),
        ).json()
        with client.stream(
            "GET",
            f"/v1/runs/{run['id']}/stream",
            headers=headers,
        ) as response:
            events = _parse_sse(response.read().decode("utf-8"))
        diagnostics = client.get(
            f"/v1/runs/{run['id']}/diagnostics",
            headers=headers,
        ).json()

    task_receipts = [
        receipt for receipt in diagnostics["tool_receipts"] if receipt["tool_name"] == "task"
    ]
    assert [receipt["status"] for receipt in task_receipts] == [
        "completed",
        "failed",
        "completed",
    ]
    assert "run.completed" in {name for name, _payload in events}
    assert "Tool call limit exceeded" not in json.dumps(events)


def test_capability_2c_subagent_tools_share_review_and_receipt_boundary(
    monkeypatch,
    tmp_path,
) -> None:
    handler = RecordingHandler(
        scripts=[
            [
                (
                    "llm.tool_call",
                    {
                        "id": "call_task",
                        "name": "task",
                        "arguments": {
                            "subagent_type": "general-purpose",
                            "description": "Write child.txt with child content.",
                        },
                    },
                ),
                ("llm.done", {"request_id": "r1", "finish_reason": "tool_calls"}),
            ],
            [
                (
                    "llm.tool_call",
                    {
                        "id": "call_child_write",
                        "name": "write_file",
                        "arguments": {"file_path": "child.txt", "content": "child"},
                    },
                ),
                ("llm.done", {"request_id": "r2", "finish_reason": "tool_calls"}),
            ],
            [
                ("llm.delta", {"content_delta": "Child file written."}),
                ("llm.done", {"request_id": "r3", "finish_reason": "stop"}),
            ],
            [
                ("llm.delta", {"content_delta": "Done."}),
                ("llm.done", {"request_id": "r4", "finish_reason": "stop"}),
            ],
        ]
    )
    with _make_client(monkeypatch, handler) as client:
        headers = {"Authorization": "Bearer tok"}
        workspace = client.post(
            "/v1/workspaces",
            headers=headers,
            json={"path": str(tmp_path), "label": "subagent"},
        )
        assert workspace.status_code == 200, workspace.text
        command = run_command("delegate a write")
        command["workspace_path"] = str(tmp_path)
        run = client.post("/v1/runs", headers=headers, json=command).json()
        with client.stream("GET", f"/v1/runs/{run['id']}/stream", headers=headers) as response:
            first_body = response.read().decode("utf-8")
        first_events = _parse_sse(first_body)
        permission = next(event[1] for event in first_events if event[0] == "permission.required")
        assert permission["tool_call_id"] == "call_child_write"
        approved = client.post(
            f"/v1/permissions/{permission['request_id']}",
            headers=headers,
            json={"decision": "approve", "scope": "once"},
        )
        assert approved.status_code == 200, approved.text
        with client.stream("GET", f"/v1/runs/{run['id']}/stream", headers=headers) as response:
            second_body = response.read().decode("utf-8")
        diagnostics = client.get(f"/v1/runs/{run['id']}/diagnostics", headers=headers).json()

    assert "run.completed" in [event[0] for event in _parse_sse(second_body)]
    assert (tmp_path / "child.txt").read_text(encoding="utf-8") == "child"
    assert {receipt["tool_call_id"] for receipt in diagnostics["tool_receipts"]} >= {
        "call_task",
        "call_child_write",
    }


def test_capability_2d_team_graph_handoff_uses_nested_task_receipts(monkeypatch) -> None:
    handler = RecordingHandler(
        scripts=[
            [
                (
                    "llm.tool_call",
                    {
                        "id": "call_team",
                        "name": "team.run",
                        "arguments": {
                            "objective": "Research evidence, then review it.",
                            "assignments": [
                                {
                                    "id": "research",
                                    "member": "researcher",
                                    "task": "Collect one evidence summary.",
                                    "output_kind": "finding",
                                },
                                {
                                    "id": "review",
                                    "member": "writer",
                                    "task": "Review the evidence summary.",
                                    "output_kind": "review",
                                    "depends_on": ["research"],
                                    "handoff_from": "research",
                                },
                            ],
                        },
                    },
                ),
                ("llm.done", {"request_id": "team-parent-1", "finish_reason": "tool_calls"}),
            ],
            [
                ("llm.delta", {"content_delta": "Evidence collected."}),
                ("llm.done", {"request_id": "team-research", "finish_reason": "stop"}),
            ],
            [
                ("llm.delta", {"content_delta": "Evidence reviewed."}),
                ("llm.done", {"request_id": "team-review", "finish_reason": "stop"}),
            ],
            [
                ("llm.delta", {"content_delta": "Team workflow complete."}),
                ("llm.done", {"request_id": "team-parent-2", "finish_reason": "stop"}),
            ],
        ]
    )
    with _make_client(monkeypatch, handler) as client:
        headers = {"Authorization": "Bearer tok"}
        run = client.post(
            "/v1/runs",
            headers=headers,
            json=run_command("run a research and review team"),
        ).json()
        with client.stream("GET", f"/v1/runs/{run['id']}/stream", headers=headers) as response:
            events = _parse_sse(response.read().decode("utf-8"))
        diagnostics = client.get(f"/v1/runs/{run['id']}/diagnostics", headers=headers).json()

    assert "run.completed" in [event[0] for event in events]
    lifecycle = [event for event in events if event[0].startswith("subagent.")]
    assert [event[0] for event in lifecycle] == [
        "subagent.spawned",
        "subagent.started",
        "subagent.completed",
        "subagent.spawned",
        "subagent.started",
        "subagent.completed",
    ]
    receipts = diagnostics["tool_receipts"]
    team_receipt = next(item for item in receipts if item["tool_name"] == "team.run")
    task_receipts = [item for item in receipts if item["tool_name"] == "task"]
    assert len(task_receipts) == 2
    assert all(
        event[1]["parent_operation_id"] == team_receipt["operation_id"] for event in lifecycle
    )
    writer_request = handler.requests[2]
    writer_user_text = " ".join(
        str(message.get("content") or "")
        for message in writer_request.get("messages", [])
        if message.get("role") == "user"
    )
    assert "Evidence collected." in writer_user_text
    assert "summaries and Artifact references only" in writer_user_text


def test_capability_2e_durable_child_spawn_wait_and_parent_projection(
    monkeypatch,
) -> None:
    class DurableChildHandler(RecordingHandler):
        def __init__(self) -> None:
            super().__init__(scripts=[])
            self._lock = threading.Lock()
            self._parent_round = 0

        def __call__(self, request: httpx.Request) -> httpx.Response:
            body = json.loads(request.read())
            with self._lock:
                self.requests.append(body)
            body_text = json.dumps(body, ensure_ascii=False)
            if "P9 final-answer reviewer" in body_text:
                return _sse(
                    [
                        (
                            "llm.delta",
                            {
                                "content_delta": json.dumps(
                                    {"decision": "allow", "reason": "complete"}
                                )
                            },
                        ),
                        (
                            "llm.done",
                            {"request_id": "completion-review", "finish_reason": "stop"},
                        ),
                    ]
                )
            if "conversation title generator" in body_text:
                return _sse(
                    [
                        ("llm.delta", {"content_delta": "Durable child run"}),
                        ("llm.done", {"request_id": "title", "finish_reason": "stop"}),
                    ]
                )
            if "<agent-role>" in body_text:
                return _sse(
                    [
                        ("llm.delta", {"content_delta": "Child evidence complete."}),
                        ("llm.done", {"request_id": "child", "finish_reason": "stop"}),
                    ]
                )

            with self._lock:
                self._parent_round += 1
                parent_round = self._parent_round
            if parent_round == 1:
                return _sse(
                    [
                        (
                            "llm.tool_call",
                            {
                                "id": "call_child_spawn",
                                "name": "child.spawn",
                                "arguments": {
                                    "agent": "researcher",
                                    "task": "Collect one primary-source finding.",
                                },
                            },
                        ),
                        ("llm.done", {"request_id": "parent-1", "finish_reason": "tool_calls"}),
                    ]
                )
            if parent_round == 2:
                child_id = ""
                for message in reversed(body.get("messages", [])):
                    content = message.get("content") if isinstance(message, dict) else None
                    if not isinstance(content, str):
                        continue
                    try:
                        candidate = json.loads(content)
                    except json.JSONDecodeError:
                        continue
                    if isinstance(candidate, dict) and candidate.get("run_kind") == "child":
                        child_id = str(candidate.get("id") or "")
                        break
                assert child_id
                return _sse(
                    [
                        (
                            "llm.tool_call",
                            {
                                "id": "call_child_wait",
                                "name": "child.wait",
                                "arguments": {
                                    "run_ids": [child_id],
                                    "condition": "all",
                                    "timeout_seconds": 5,
                                },
                            },
                        ),
                        ("llm.done", {"request_id": "parent-2", "finish_reason": "tool_calls"}),
                    ]
                )
            return _sse(
                [
                    ("llm.delta", {"content_delta": "Parent collected the child result."}),
                    ("llm.done", {"request_id": "parent-3", "finish_reason": "stop"}),
                ]
            )

    handler = DurableChildHandler()
    with _make_client(monkeypatch, handler) as client:
        headers = {"Authorization": "Bearer tok"}
        client.app.state.coordinator._slots = asyncio.Semaphore(1)
        run = client.post(
            "/v1/runs",
            headers=headers,
            json=run_command("delegate durable research"),
        ).json()
        with client.stream("GET", f"/v1/runs/{run['id']}/stream", headers=headers) as response:
            events = _parse_sse(response.read().decode("utf-8"))
        children = client.get(f"/v1/runs/{run['id']}/children", headers=headers).json()["children"]
        child = client.get(f"/v1/runs/{children[0]['id']}", headers=headers).json()

    assert [event[0] for event in events if event[0].startswith("child.")] == [
        "child.spawned",
        "child.started",
        "child.completed",
    ]
    assert len(children) == 1
    assert children[0]["status"] == "completed"
    assert children[0]["result"] == "Child evidence complete."
    assert child["root_run_id"] == run["root_run_id"]
    assert child["parent_run_id"] == run["id"]


def test_capability_2g_parent_settlement_waits_for_required_child_without_model_wait(
    monkeypatch,
) -> None:
    class CoordinatedChildHandler(RecordingHandler):
        def __init__(self) -> None:
            super().__init__(scripts=[])
            self._lock = threading.Lock()
            self._parent_round = 0

        def __call__(self, request: httpx.Request) -> httpx.Response:
            body = json.loads(request.read())
            with self._lock:
                self.requests.append(body)
            body_text = json.dumps(body, ensure_ascii=False)
            if "P9 final-answer reviewer" in body_text:
                return _sse(
                    [
                        (
                            "llm.delta",
                            {
                                "content_delta": json.dumps(
                                    {"decision": "allow", "reason": "complete"}
                                )
                            },
                        ),
                        ("llm.done", {"request_id": "review", "finish_reason": "stop"}),
                    ]
                )
            if "conversation title generator" in body_text:
                return _sse(
                    [
                        ("llm.delta", {"content_delta": "Coordinated child"}),
                        ("llm.done", {"request_id": "title", "finish_reason": "stop"}),
                    ]
                )
            if "<agent-role>" in body_text:
                return _sse(
                    [
                        ("llm.delta", {"content_delta": "Required child finished."}),
                        ("llm.done", {"request_id": "child", "finish_reason": "stop"}),
                    ]
                )

            with self._lock:
                self._parent_round += 1
                parent_round = self._parent_round
            if parent_round == 1:
                return _sse(
                    [
                        (
                            "llm.tool_call",
                            {
                                "id": "call_required_child",
                                "name": "child.spawn",
                                "arguments": {
                                    "agent": "researcher",
                                    "task": "Finish before the parent commits.",
                                    "completion_mode": "required",
                                },
                            },
                        ),
                        ("llm.done", {"request_id": "parent-1", "finish_reason": "tool_calls"}),
                    ]
                )
            return _sse(
                [
                    ("llm.delta", {"content_delta": "Parent answer proposed early."}),
                    ("llm.done", {"request_id": "parent-2", "finish_reason": "stop"}),
                ]
            )

    handler = CoordinatedChildHandler()
    with _make_client(monkeypatch, handler) as client:
        headers = {"Authorization": "Bearer tok"}
        client.app.state.coordinator._slots = asyncio.Semaphore(1)
        run = client.post(
            "/v1/runs",
            headers=headers,
            json=run_command("complete only after required child"),
        ).json()
        with client.stream("GET", f"/v1/runs/{run['id']}/stream", headers=headers) as response:
            events = _parse_sse(response.read().decode("utf-8"))
        snapshot = client.get(
            f"/v1/runs/{run['id']}/collaboration",
            headers=headers,
        ).json()

    durable_names = [
        event[0] for event in events if event[0].startswith("child.") or event[0] == "run.completed"
    ]
    assert durable_names[-2:] == ["child.completed", "run.completed"]
    assert snapshot["root"]["status"] == "completed"
    assert snapshot["children"][0]["status"] == "completed"
    assert snapshot["children"][0]["completion_mode"] == "required"
    assert snapshot["completion"]["satisfied"] is True
    assert snapshot["completion"]["wait_for"] == []


def test_capability_2f_root_and_durable_child_exchange_acknowledged_mailbox_messages(
    monkeypatch,
) -> None:
    class MailboxHandler(RecordingHandler):
        def __init__(self) -> None:
            super().__init__(scripts=[])
            self._lock = threading.Lock()
            self._parent_round = 0
            self._child_poll = 0
            self._child_replied = False
            self._parent_acked = False

        @staticmethod
        def mailbox_message(body: dict[str, Any]) -> dict[str, Any] | None:
            for message in reversed(body.get("messages", [])):
                content = message.get("content") if isinstance(message, dict) else None
                if not isinstance(content, str) or "【同一协作任务中的 Agent 消息】" not in content:
                    continue
                try:
                    return json.loads(content.rsplit("\n\n", 1)[-1])
                except json.JSONDecodeError:
                    return None
            return None

        @staticmethod
        def child_run_id(body: dict[str, Any]) -> str:
            for message in reversed(body.get("messages", [])):
                content = message.get("content") if isinstance(message, dict) else None
                if not isinstance(content, str) or not content.startswith("{"):
                    continue
                try:
                    candidate = json.loads(content)
                except json.JSONDecodeError:
                    continue
                if isinstance(candidate, dict) and candidate.get("run_kind") == "child":
                    return str(candidate.get("id") or "")
            return ""

        def __call__(self, request: httpx.Request) -> httpx.Response:
            body = json.loads(request.read())
            with self._lock:
                self.requests.append(body)
            body_text = json.dumps(body, ensure_ascii=False)
            if "P9 final-answer reviewer" in body_text:
                return _sse(
                    [
                        (
                            "llm.delta",
                            {
                                "content_delta": json.dumps(
                                    {"decision": "allow", "reason": "complete"}
                                )
                            },
                        ),
                        ("llm.done", {"request_id": "review", "finish_reason": "stop"}),
                    ]
                )
            if "conversation title generator" in body_text:
                return _sse(
                    [
                        ("llm.delta", {"content_delta": "Agent mailbox"}),
                        ("llm.done", {"request_id": "title", "finish_reason": "stop"}),
                    ]
                )
            if "<agent-role>" in body_text:
                incoming = self.mailbox_message(body)
                if incoming is None:
                    with self._lock:
                        self._child_poll += 1
                        poll = self._child_poll
                    return _sse(
                        [
                            (
                                "llm.tool_call",
                                {
                                    "id": f"call_child_inbox_{poll}",
                                    "name": "mailbox.inbox",
                                    "arguments": {},
                                },
                            ),
                            (
                                "llm.done",
                                {"request_id": f"child-poll-{poll}", "finish_reason": "tool_calls"},
                            ),
                        ]
                    )
                with self._lock:
                    should_reply = not self._child_replied
                    self._child_replied = True
                if should_reply:
                    return _sse(
                        [
                            (
                                "llm.tool_call",
                                {
                                    "id": "call_child_reply",
                                    "name": "mailbox.reply",
                                    "arguments": {
                                        "in_reply_to": incoming["message_id"],
                                        "kind": "result",
                                        "text": "Primary source confirmed.",
                                    },
                                },
                            ),
                            (
                                "llm.tool_call",
                                {
                                    "id": "call_child_ack",
                                    "name": "mailbox.ack",
                                    "arguments": {"message_ids": [incoming["message_id"]]},
                                },
                            ),
                            (
                                "llm.done",
                                {"request_id": "child-1", "finish_reason": "tool_calls"},
                            ),
                        ]
                    )
                return _sse(
                    [
                        ("llm.delta", {"content_delta": "Child mailbox work complete."}),
                        ("llm.done", {"request_id": "child-2", "finish_reason": "stop"}),
                    ]
                )

            with self._lock:
                self._parent_round += 1
                parent_round = self._parent_round
            if parent_round == 1:
                return _sse(
                    [
                        (
                            "llm.tool_call",
                            {
                                "id": "call_mailbox_spawn",
                                "name": "child.spawn",
                                "arguments": {
                                    "agent": "researcher",
                                    "task": "Confirm the primary source.",
                                },
                            },
                        ),
                        ("llm.done", {"request_id": "parent-1", "finish_reason": "tool_calls"}),
                    ]
                )
            child_id = self.child_run_id(body)
            if parent_round == 2:
                assert child_id
                return _sse(
                    [
                        (
                            "llm.tool_call",
                            {
                                "id": "call_parent_send",
                                "name": "mailbox.send",
                                "arguments": {
                                    "recipient_run_id": child_id,
                                    "kind": "request",
                                    "text": "Return the verified source result.",
                                },
                            },
                        ),
                        ("llm.done", {"request_id": "parent-2", "finish_reason": "tool_calls"}),
                    ]
                )
            if parent_round == 3:
                assert child_id
                return _sse(
                    [
                        (
                            "llm.tool_call",
                            {
                                "id": "call_parent_wait",
                                "name": "child.wait",
                                "arguments": {
                                    "run_ids": [child_id],
                                    "condition": "all",
                                    "timeout_seconds": 5,
                                },
                            },
                        ),
                        ("llm.done", {"request_id": "parent-3", "finish_reason": "tool_calls"}),
                    ]
                )
            incoming = self.mailbox_message(body)
            with self._lock:
                should_ack = incoming is not None and not self._parent_acked
                if should_ack:
                    self._parent_acked = True
            if should_ack:
                return _sse(
                    [
                        (
                            "llm.tool_call",
                            {
                                "id": "call_parent_ack",
                                "name": "mailbox.ack",
                                "arguments": {"message_ids": [incoming["message_id"]]},
                            },
                        ),
                        ("llm.done", {"request_id": "parent-4", "finish_reason": "tool_calls"}),
                    ]
                )
            return _sse(
                [
                    ("llm.delta", {"content_delta": "Parent received the verified result."}),
                    ("llm.done", {"request_id": "parent-5", "finish_reason": "stop"}),
                ]
            )

    handler = MailboxHandler()
    with _make_client(monkeypatch, handler) as client:
        headers = {"Authorization": "Bearer tok"}
        client.app.state.coordinator._slots = asyncio.Semaphore(1)
        run = client.post(
            "/v1/runs",
            headers=headers,
            json=run_command("coordinate through a durable mailbox"),
        ).json()
        with client.stream("GET", f"/v1/runs/{run['id']}/stream", headers=headers) as response:
            events = _parse_sse(response.read().decode("utf-8"))
        child = client.get(f"/v1/runs/{run['id']}/children", headers=headers).json()["children"][0]
        parent_inbox = client.get(
            f"/v1/runs/{run['id']}/mailbox?box=inbox", headers=headers
        ).json()["messages"]
        child_inbox = client.get(
            f"/v1/runs/{child['id']}/mailbox?box=inbox", headers=headers
        ).json()["messages"]

    assert [message["status"] for message in parent_inbox] == ["acknowledged"]
    assert [message["status"] for message in child_inbox] == ["acknowledged"]
    assert parent_inbox[0]["text"] == "Primary source confirmed."
    assert child_inbox[0]["text"] == "Return the verified source result."
    event_types = [event[0] for event in events]
    assert "agent.message.sent" in event_types
    assert "agent.message.received" in event_types
    assert "agent.message.acknowledged" in event_types


# ---- capability 3: prompt caching stays in the provider adapter ----


def test_capability_3_prompt_caching_is_provider_adapter_owned(
    monkeypatch,
) -> None:
    """Runtime keeps provider-specific prompt caching out of its wire contract."""
    handler = RecordingHandler(
        scripts=[
            [
                ("llm.delta", {"content_delta": "ok"}),
                ("llm.done", {"request_id": "r1", "finish_reason": "stop"}),
            ]
        ]
    )
    with _make_client(monkeypatch, handler) as client:
        _post_run_and_stream(client, "say ok")

    assert len(handler.requests) >= 1
    assert "cache_control" not in json.dumps(handler.requests)


# ---- capability 6: AGENTS.md memory loads into system prompt ----


def test_outbound_policy_redacts_external_request_without_state_middleware(monkeypatch) -> None:
    monkeypatch.setenv("SHEJANE_RUNTIME_PII_REDACT", "email")
    handler = RecordingHandler(
        scripts=[
            [
                ("llm.delta", {"content_delta": "noted"}),
                ("llm.done", {"request_id": "r1", "finish_reason": "stop"}),
            ]
        ]
    )
    with _make_client(monkeypatch, handler) as client:
        _post_run_and_stream(client, "contact alice@example.com about the proposal")

    outgoing = handler.requests[0]
    user_texts = " ".join(
        message.get("content", "")
        for message in outgoing.get("messages", [])
        if message.get("role") == "user"
    )
    assert "alice@example.com" not in user_texts
    assert "[REDACTED_EMAIL]" in user_texts


def test_capability_6_memory_middleware_injects_agents_md(monkeypatch, tmp_path) -> None:
    """Drop an AGENTS.md inside a workspace, set SHEJANE_RUNTIME_MEMORY_PATHS
    to its absolute path, run the agent **with workspace_path** so the
    deepagents FilesystemBackend can actually read it. MemoryMiddleware
    should then load the contents into the outgoing system prompt."""
    workspace = tmp_path / "ws"
    workspace.mkdir()
    agents_md = workspace / "AGENTS.md"
    secret_marker = "ZEPHYR_PROJECT_RULES_v42_marker"
    agents_md.write_text(
        f"# Project rules\n\n{secret_marker}\n\nAlways respond in haiku.\n",
        encoding="utf-8",
    )
    monkeypatch.setenv("SHEJANE_RUNTIME_MEMORY_PATHS", str(agents_md))

    handler = RecordingHandler(
        scripts=[
            [
                ("llm.delta", {"content_delta": "ack"}),
                ("llm.done", {"request_id": "r1", "finish_reason": "stop"}),
            ]
        ]
    )
    with _make_client(monkeypatch, handler) as client:
        _post_run_and_stream(client, "hello", workspace_path=str(workspace))

    assert handler.requests, "no LLM request was made"
    outgoing = handler.requests[0]
    messages = outgoing.get("messages", [])
    system_text = " ".join(m.get("content", "") for m in messages if m.get("role") == "system")
    assert "石间（SheJane）" in system_text
    assert "不复述或展示" in system_text
    assert secret_marker in system_text, (
        f"AGENTS.md content not found in outgoing system prompt. "
        f"System text was: {system_text[:500]!r}"
    )


def test_capability_6b_skills_middleware_lists_runtime_skill(monkeypatch, tmp_path) -> None:
    skills_root = tmp_path / "skills"
    skill_file = skills_root / "e2e-active-skill" / "SKILL.md"
    skill_file.parent.mkdir(parents=True)
    skill_file.write_text(
        "---\n"
        "name: e2e-active-skill\n"
        "description: Runtime Skill prompt probe.\n"
        "---\n"
        "\nReply with E2E_SKILL_ACTIVE.\n",
        encoding="utf-8",
    )
    monkeypatch.setenv("SHEJANE_RUNTIME_SKILLS_PATH", str(skills_root))
    handler = RecordingHandler(
        scripts=[
            [
                ("llm.delta", {"content_delta": "ack"}),
                ("llm.done", {"request_id": "r1", "finish_reason": "stop"}),
            ]
        ]
    )

    with _make_client(monkeypatch, handler) as client:
        _post_run_and_stream(client, "Use the active skill", settings={"skills": "on"})

    system_text = " ".join(
        message.get("content", "")
        for message in handler.requests[0].get("messages", [])
        if message.get("role") == "system"
    )
    assert "e2e-active-skill" in system_text
    assert str(skill_file) in system_text


def test_weather_lookup_reaches_model_with_only_relevant_tool_families(monkeypatch) -> None:
    handler = RecordingHandler(
        scripts=[
            [
                ("llm.delta", {"content_delta": "杭州今天多云，出门带伞。"}),
                ("llm.done", {"request_id": "weather", "finish_reason": "stop"}),
            ]
        ]
    )

    with _make_client(monkeypatch, handler) as client:
        events = _post_run_and_stream(client, "帮我查一下 今天杭州的天气")

    tool_names = {
        str(tool.get("name") or tool.get("function", {}).get("name"))
        for tool in handler.requests[0]["tools"]
    }
    assert "run.completed" in {name for name, _payload in events}
    assert "run.failed" not in {name for name, _payload in events}
    assert "web_fetch" in tool_names
    assert tool_names.isdisjoint(
        {"read_file", "ls", "glob", "execute", "task", "child_spawn", "child_check"}
    )
    assert handler.agent_requests == 1


def test_weather_lookup_fetches_data_and_completes_without_summarization(monkeypatch) -> None:
    from shejane_runtime.tools import web as web_module

    weather_url = "https://weather.example.test/hangzhou/today"
    weather_requests: list[str] = []

    def weather_handler(request: httpx.Request) -> httpx.Response:
        weather_requests.append(str(request.url))
        return httpx.Response(
            200,
            json={"city": "杭州", "condition": "多云", "temperature_c": 27},
            request=request,
        )

    monkeypatch.setattr(
        web_module,
        "_pinned_transport",
        lambda _url: (httpx.MockTransport(weather_handler), ""),
    )
    handler = RecordingHandler(
        scripts=[
            [
                (
                    "llm.tool_call",
                    {
                        "id": "weather-fetch",
                        "name": "web.fetch",
                        "arguments": {"url": weather_url},
                    },
                ),
                ("llm.done", {"request_id": "weather-1", "finish_reason": "tool_calls"}),
            ],
            [
                ("llm.delta", {"content_delta": "杭州今天 27°C，多云。"}),
                ("llm.done", {"request_id": "weather-2", "finish_reason": "stop"}),
            ],
        ]
    )

    with _make_client(monkeypatch, handler) as client:
        headers = {"Authorization": "Bearer tok"}
        run = client.post(
            "/v1/runs",
            headers=headers,
            json=run_command("帮我查一下 今天杭州的天气"),
        )
        assert run.status_code == 200, run.text
        run_id = run.json()["id"]
        with client.stream("GET", f"/v1/runs/{run_id}/stream", headers=headers) as response:
            events = _parse_sse(response.read().decode("utf-8"))
        diagnostics = client.get(f"/v1/runs/{run_id}/diagnostics", headers=headers).json()

    completed = next(payload for name, payload in events if name == "run.completed")
    tool_result = next(payload for name, payload in events if name == "tool.completed")
    weather = json.loads(json.loads(tool_result["content"])["body"])
    model_spans = [span for span in diagnostics["trace"]["spans"] if span["kind"] == "model"]
    assert weather_requests == [weather_url]
    assert weather == {"city": "杭州", "condition": "多云", "temperature_c": 27}
    assert completed["final_text"] == "杭州今天 27°C，多云。"
    assert [span["name"] for span in model_spans].count("agent") == 2
    assert "summarization" not in {span["name"] for span in model_spans}
    assert diagnostics["tool_receipts"][0]["status"] == "completed"
    assert handler.agent_requests == 2


# ---- capability 7: TodoList middleware exposes write_todos ----


def test_capability_7_todolist_middleware_exposes_write_todos_tool(monkeypatch, tmp_path) -> None:
    """write_todos should appear in the compiled agent's tool registry."""
    from shejane_runtime.agent.builder import build_agent, open_checkpointer
    from shejane_runtime.store.sqlite import LocalStore

    async def run() -> set[str]:
        reset_settings_for_tests(data_dir=tmp_path)
        monkeypatch.delenv("SHEJANE_RUNTIME_MCP_SERVERS", raising=False)
        monkeypatch.delenv("TAVILY_API_KEY", raising=False)
        store = await LocalStore.open(tmp_path / "store.db")
        saver, stack = await open_checkpointer()
        try:
            agent = await build_agent(
                store=store,
                checkpointer=saver,
                workspace_root=str(tmp_path),
                run_id="t-todo-1",
            )
            tools_node = agent.nodes.get("tools")
            if tools_node is None:
                return set()
            bound = getattr(tools_node, "bound", None)
            return set(getattr(bound, "tools_by_name", {}).keys())
        finally:
            await store.close()
            await stack.aclose()

    names = asyncio.run(run())
    assert "write_todos" in names, f"write_todos missing. tools: {sorted(names)}"


# ---- bonus: a "happy path" capability sanity (capability 8) ----


def test_capability_2b_subagent_parallel_dispatch(monkeypatch) -> None:
    """Verify the LLM can dispatch **multiple** task() subagents in one
    turn and merge their results without leaking private control state."""
    handler = RecordingHandler(
        scripts=[
            [
                ("llm.delta", {"content_delta": "Dispatching two researchers. "}),
                (
                    "llm.tool_call",
                    {
                        "id": "call_p1",
                        "name": "task",
                        "arguments": {
                            "subagent_type": "researcher",
                            "description": "subquery A: LangGraph 1.x changes",
                        },
                    },
                ),
                (
                    "llm.tool_call",
                    {
                        "id": "call_p2",
                        "name": "task",
                        "arguments": {
                            "subagent_type": "researcher",
                            "description": "subquery B: deepagents adoption",
                        },
                    },
                ),
                ("llm.done", {"request_id": "rp", "finish_reason": "tool_calls"}),
            ],
            # Both researcher subagents share these scripted responses
            # (RecordingHandler pops them in order — either order works
            # because we just assert both spawn events are present).
            [
                ("llm.delta", {"content_delta": "found A"}),
                ("llm.done", {"request_id": "ra", "finish_reason": "stop"}),
            ],
            [
                ("llm.delta", {"content_delta": "found B"}),
                ("llm.done", {"request_id": "rb", "finish_reason": "stop"}),
            ],
            [
                ("llm.delta", {"content_delta": "Combined findings."}),
                ("llm.done", {"request_id": "final", "finish_reason": "stop"}),
            ],
        ]
    )
    with _make_client(monkeypatch, handler) as client:
        events = _post_run_and_stream(client, "compare two angles of LangGraph adoption")

    spawn_events = [e for e in events if e[0] == "subagent.spawned"]
    assert len(spawn_events) >= 2, (
        f"expected at least 2 subagent.spawned events. got {len(spawn_events)}: "
        f"{[e[1].get('id') for e in spawn_events]}"
    )
    operation_ids = {e[1].get("operation_id") for e in spawn_events}
    assert len(operation_ids) == 2
    assert {e[1].get("tool_call_id") for e in spawn_events} == {"call_p1", "call_p2"}
    assert "run.completed" in {name for name, _payload in events}
    assert "run.failed" not in {name for name, _payload in events}


def test_capability_9_memory_search_tool_in_agent(monkeypatch, tmp_path) -> None:
    """`memory.search` tool must appear in the compiled agent's toolset —
    that's the read-side of the long-term memory loop."""
    from shejane_runtime.agent.builder import build_agent, open_checkpointer, open_store
    from shejane_runtime.store.sqlite import LocalStore

    async def run() -> set[str]:
        reset_settings_for_tests(data_dir=tmp_path)
        monkeypatch.delenv("SHEJANE_RUNTIME_MCP_SERVERS", raising=False)
        monkeypatch.delenv("TAVILY_API_KEY", raising=False)
        store = await LocalStore.open(tmp_path / "store.db")
        saver, ck_stack = await open_checkpointer()
        agent_store, st_stack = await open_store()
        try:
            agent = await build_agent(
                store=store,
                checkpointer=saver,
                agent_store=agent_store,
                workspace_root=str(tmp_path),
                run_id="t-mem-1",
            )
            tools_node = agent.nodes.get("tools")
            if tools_node is None:
                return set()
            bound = getattr(tools_node, "bound", None)
            return set(getattr(bound, "tools_by_name", {}).keys())
        finally:
            await store.close()
            await ck_stack.aclose()
            await st_stack.aclose()

    names = asyncio.run(run())
    assert "memory.search" in names


def test_capability_10_plan_first_is_a_runtime_gate_not_prompt_text(monkeypatch) -> None:
    """Complex work is repaired into sequential todo transitions at P9."""
    monkeypatch.setenv("SHEJANE_PLAN_FIRST", "always")
    handler = RecordingHandler(
        scripts=[
            [
                (
                    "llm.tool_call",
                    {
                        "id": "too-early",
                        "name": "write_file",
                        "arguments": {"file_path": "too-early.txt", "content": "no"},
                    },
                ),
                ("llm.done", {"request_id": "r1", "finish_reason": "tool_calls"}),
            ],
            [
                (
                    "llm.tool_call",
                    {
                        "id": "plan-start",
                        "name": "write_todos",
                        "arguments": {
                            "todos": [
                                {"content": "First slice", "status": "in_progress"},
                                {"content": "Second slice", "status": "pending"},
                            ]
                        },
                    },
                ),
                ("llm.done", {"request_id": "r2", "finish_reason": "tool_calls"}),
            ],
            [
                (
                    "llm.tool_call",
                    {
                        "id": "plan-advance",
                        "name": "write_todos",
                        "arguments": {
                            "todos": [
                                {"content": "First slice", "status": "completed"},
                                {"content": "Second slice", "status": "in_progress"},
                            ]
                        },
                    },
                ),
                ("llm.done", {"request_id": "r3", "finish_reason": "tool_calls"}),
            ],
            [
                (
                    "llm.tool_call",
                    {
                        "id": "plan-finish",
                        "name": "write_todos",
                        "arguments": {
                            "todos": [
                                {"content": "First slice", "status": "completed"},
                                {"content": "Second slice", "status": "completed"},
                            ]
                        },
                    },
                ),
                ("llm.done", {"request_id": "r4", "finish_reason": "tool_calls"}),
            ],
            [
                ("llm.delta", {"content_delta": "incremental flow complete"}),
                ("llm.done", {"request_id": "r5", "finish_reason": "stop"}),
            ],
        ]
    )
    with _make_client(monkeypatch, handler) as client:
        events = _post_run_and_stream(client, "do the thing")

    first_system = " ".join(
        str(message.get("content") or "")
        for message in handler.requests[0].get("messages", [])
        if message.get("role") == "system"
    )
    assert "Plan-First protocol" not in first_system
    assert "Call write_todos before any work tool" in str(handler.requests[1])
    assert any(name == "run.completed" for name, _payload in events)


def test_capability_8_happy_path_run_completes(monkeypatch) -> None:
    """End-to-end: a clean run goes from POST → SSE → run.completed."""
    handler = RecordingHandler(
        scripts=[
            [
                ("llm.delta", {"content_delta": "All done. "}),
                ("llm.done", {"request_id": "r1", "finish_reason": "stop"}),
            ]
        ]
    )
    with _make_client(monkeypatch, handler) as client:
        events = _post_run_and_stream(client, "say hi")

    event_names = [e[0] for e in events]
    assert "run.started" in event_names
    assert "run.completed" in event_names
    for name, data in events:
        if name == "run.completed":
            final = data.get("final_text", "") if isinstance(data, dict) else ""
            assert "All done" in final
            break
