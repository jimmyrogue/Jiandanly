"""Word outline and copy-on-first-write editing tools."""

from __future__ import annotations

from copy import deepcopy
from typing import Any

from docx import Document as DocxDocument
from docx.document import Document as DocxDocumentType
from docx.text.paragraph import Paragraph as DocxParagraph
from langchain_core.tools import tool

from .office_common import (
    atomic_write as _atomic_write,
)
from .office_common import (
    ensure_copy_for_write as _ensure_copy_for_write,
)
from .office_common import (
    validate_path as _validate_path,
)
from .office_common import (
    write_error as _write_error,
)
from .office_common import (
    write_result as _write_result,
)


def outline_docx(path: Any) -> dict[str, Any]:
    """Return heading, paragraph, and table metadata for a Word document."""
    doc = DocxDocument(path)
    headings: list[dict[str, Any]] = []
    paragraph_count = 0
    table_count = len(doc.tables)
    for para in doc.paragraphs:
        paragraph_count += 1
        style_name = (para.style.name if para.style is not None else "") or ""
        if style_name.startswith("Heading"):
            level_str = style_name.split()[-1] if " " in style_name else "1"
            try:
                level = int(level_str)
            except ValueError:
                level = 1
            text = para.text.strip()
            if text:
                headings.append({"level": level, "text": text})
    return {
        "headings": headings,
        "paragraph_count": paragraph_count,
        "table_count": table_count,
    }


def _find_paragraph(doc: DocxDocumentType, target: str) -> DocxParagraph | None:
    """Return the first paragraph containing target, including table cells."""
    if not target:
        return None
    for para in doc.paragraphs:
        if target in para.text:
            return para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if target in para.text:
                        return para
    return None


def _set_paragraph_text(para: DocxParagraph, new_text: str) -> None:
    """Replace paragraph text while preserving its paragraph-level style."""
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(new_text)


def _insert_paragraph_relative(
    anchor: DocxParagraph, content: str, position: str, style: str | None
) -> DocxParagraph:
    """Insert a paragraph immediately before or after anchor."""
    new_p = deepcopy(anchor._p)
    for child in list(new_p):
        if child.tag.endswith("}r"):
            new_p.remove(child)
    if position == "before":
        anchor._p.addprevious(new_p)
    else:
        anchor._p.addnext(new_p)
    new_para = DocxParagraph(new_p, anchor._parent)
    if style:
        new_para.style = style
    new_para.add_run(content)
    return new_para


def _iter_all_paragraphs(doc: DocxDocumentType):
    """Yield body paragraphs and paragraphs nested in table cells."""
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


@tool("office.find_replace")
def office_find_replace(
    path: str,
    find: str,
    replace: str,
    count: int | None = None,
) -> dict[str, Any]:
    """Replace text across a .docx. Writes to a copy named `<basename>.edited.docx`.

    The original file is NEVER modified — you always operate on a copy.
    If the copy already exists from a previous edit, this writes to
    that same copy (chained edits land in one file).

    Args:
        path: .docx to edit. May be the original or its `.edited` copy
              (idempotent: if you pass the edited copy, we write back
              to it).
        find: text to search for (exact substring match).
        replace: text to substitute in.
        count: optional — stop after this many replacements. None =
               replace all.

    Returns:
        dict with `ok`, `original_path`, `edited_path`, `kind`,
        `summary={replaced}`. Use `edited_path` for subsequent reads
        or edits on this document.

    Note on formatting: this tool collapses run-level char formatting
    in each modified paragraph (paragraph-level style is preserved).
    If you need to keep "bold middle word" intact, use
    `office.update_paragraph` instead — it's whole-paragraph anyway.
    """
    resolved, kind, err = _validate_path(path)
    if err is not None:
        return _write_error(None, err)
    assert resolved is not None and kind is not None
    if kind != "word":
        return _write_error(kind, "office.find_replace requires a .docx file", resolved)
    if not find:
        return _write_error(kind, "find text required", resolved)

    target = _ensure_copy_for_write(resolved)
    remaining = count if (count is None or count > 0) else 0

    def _do_write(tmp_path: str) -> dict[str, Any]:
        nonlocal remaining
        doc = DocxDocument(target)
        replaced = 0
        for para in _iter_all_paragraphs(doc):
            if remaining == 0:
                break
            text = para.text
            if find not in text:
                continue
            occurrences = text.count(find)
            if remaining is not None:
                occurrences = min(occurrences, remaining)
            _set_paragraph_text(para, text.replace(find, replace, occurrences))
            replaced += occurrences
            if remaining is not None:
                remaining -= occurrences
        doc.save(tmp_path)
        return {"replaced": replaced}

    try:
        summary = _atomic_write(target, kind, _do_write)
    except Exception as exc:
        return _write_error(kind, f"write failed: {exc.__class__.__name__}: {exc}", resolved)
    return _write_result(resolved, target, kind, summary)


@tool("office.insert_paragraph")
def office_insert_paragraph(
    path: str,
    anchor: str,
    content: str,
    position: str = "after",
    style: str | None = None,
) -> dict[str, Any]:
    """Insert a new paragraph relative to an anchor paragraph in a .docx.

    Writes to `<basename>.edited.docx`; original preserved.

    Args:
        path: .docx to edit.
        anchor: text that uniquely identifies the anchor paragraph
                (first paragraph whose text contains this substring).
                Empty string OR the literal "__END__" inserts at the
                very end of the document.
        content: text of the new paragraph.
        position: "before" or "after" the anchor (default "after").
                  Ignored when anchor is "__END__".
        style: optional python-docx style name (e.g. "Heading 1",
               "Normal", "Quote", "List Bullet"). Use `office.outline`
               to see which styles already appear in the document.

    Returns:
        dict with `ok`, `original_path`, `edited_path`, `kind`,
        `summary={inserted_at, style?}`.
    """
    resolved, kind, err = _validate_path(path)
    if err is not None:
        return _write_error(None, err)
    assert resolved is not None and kind is not None
    if kind != "word":
        return _write_error(kind, "office.insert_paragraph requires a .docx file", resolved)
    if position not in {"before", "after"}:
        return _write_error(
            kind, f"position must be 'before' or 'after', got {position!r}", resolved
        )

    target = _ensure_copy_for_write(resolved)

    def _do_write(tmp_path: str) -> dict[str, Any]:
        doc = DocxDocument(target)
        if anchor in ("", "__END__"):
            new_para = (
                doc.add_paragraph(content, style=style) if style else doc.add_paragraph(content)
            )
            doc.save(tmp_path)
            return {"inserted_at": "end", "style": new_para.style.name if new_para.style else None}
        match = _find_paragraph(doc, anchor)
        if match is None:
            raise ValueError(f"anchor not found: {anchor!r}")
        new_para = _insert_paragraph_relative(match, content, position, style)
        doc.save(tmp_path)
        return {
            "inserted_at": f"{position} anchor",
            "anchor_excerpt": match.text[:60],
            "style": new_para.style.name if new_para.style else None,
        }

    try:
        summary = _atomic_write(target, kind, _do_write)
    except Exception as exc:
        return _write_error(kind, f"write failed: {exc.__class__.__name__}: {exc}", resolved)
    return _write_result(resolved, target, kind, summary)


@tool("office.update_paragraph")
def office_update_paragraph(
    path: str,
    target: str,
    content: str,
    style: str | None = None,
) -> dict[str, Any]:
    """Rewrite the first paragraph that contains `target` substring.

    Writes to `<basename>.edited.docx`; original preserved.

    Args:
        path: .docx to edit.
        target: substring that identifies the paragraph to update
                (first match wins; pick something unique).
        content: new full text of that paragraph (replaces ALL existing
                 text in the paragraph).
        style: optional new paragraph style name.

    Returns:
        dict with `ok`, `original_path`, `edited_path`, `kind`,
        `summary={updated_excerpt, style?}`.
    """
    resolved, kind, err = _validate_path(path)
    if err is not None:
        return _write_error(None, err)
    assert resolved is not None and kind is not None
    if kind != "word":
        return _write_error(kind, "office.update_paragraph requires a .docx file", resolved)
    if not target:
        return _write_error(kind, "target text required", resolved)

    target_path = _ensure_copy_for_write(resolved)

    def _do_write(tmp_path: str) -> dict[str, Any]:
        doc = DocxDocument(target_path)
        match = _find_paragraph(doc, target)
        if match is None:
            raise ValueError(f"target paragraph not found: {target!r}")
        old_excerpt = match.text[:60]
        _set_paragraph_text(match, content)
        if style:
            match.style = style
        doc.save(tmp_path)
        return {
            "updated_excerpt": old_excerpt,
            "style": match.style.name if match.style else None,
        }

    try:
        summary = _atomic_write(target_path, kind, _do_write)
    except Exception as exc:
        return _write_error(kind, f"write failed: {exc.__class__.__name__}: {exc}", resolved)
    return _write_result(resolved, target_path, kind, summary)


@tool("office.delete_paragraph")
def office_delete_paragraph(path: str, target: str) -> dict[str, Any]:
    """Delete the first paragraph that contains `target` substring.

    Writes to `<basename>.edited.docx`; original preserved.

    Args:
        path: .docx to edit.
        target: substring that identifies the paragraph (first match wins).

    Returns:
        dict with `ok`, `original_path`, `edited_path`, `kind`,
        `summary={deleted_excerpt}`.
    """
    resolved, kind, err = _validate_path(path)
    if err is not None:
        return _write_error(None, err)
    assert resolved is not None and kind is not None
    if kind != "word":
        return _write_error(kind, "office.delete_paragraph requires a .docx file", resolved)
    if not target:
        return _write_error(kind, "target text required", resolved)

    target_path = _ensure_copy_for_write(resolved)

    def _do_write(tmp_path: str) -> dict[str, Any]:
        doc = DocxDocument(target_path)
        match = _find_paragraph(doc, target)
        if match is None:
            raise ValueError(f"target paragraph not found: {target!r}")
        deleted_excerpt = match.text[:60]
        p_element = match._p
        p_element.getparent().remove(p_element)
        doc.save(tmp_path)
        return {"deleted_excerpt": deleted_excerpt}

    try:
        summary = _atomic_write(target_path, kind, _do_write)
    except Exception as exc:
        return _write_error(kind, f"write failed: {exc.__class__.__name__}: {exc}", resolved)
    return _write_result(resolved, target_path, kind, summary)


@tool("office.apply_style")
def office_apply_style(path: str, target: str, style: str) -> dict[str, Any]:
    """Change the paragraph style of the first paragraph matching `target`.

    Writes to `<basename>.edited.docx`; original preserved.

    Args:
        path: .docx to edit.
        target: substring identifying the paragraph (first match wins).
        style: python-docx style name. Common values: "Heading 1",
               "Heading 2", "Heading 3", "Normal", "Quote", "Title",
               "Subtitle", "List Bullet", "List Number".

    Returns:
        dict with `ok`, `original_path`, `edited_path`, `kind`,
        `summary={paragraph_excerpt, old_style, new_style}`.
    """
    resolved, kind, err = _validate_path(path)
    if err is not None:
        return _write_error(None, err)
    assert resolved is not None and kind is not None
    if kind != "word":
        return _write_error(kind, "office.apply_style requires a .docx file", resolved)
    if not target or not style:
        return _write_error(kind, "target and style are both required", resolved)

    target_path = _ensure_copy_for_write(resolved)

    def _do_write(tmp_path: str) -> dict[str, Any]:
        doc = DocxDocument(target_path)
        match = _find_paragraph(doc, target)
        if match is None:
            raise ValueError(f"target paragraph not found: {target!r}")
        old_style = match.style.name if match.style else None
        match.style = style
        doc.save(tmp_path)
        return {
            "paragraph_excerpt": match.text[:60],
            "old_style": old_style,
            "new_style": match.style.name if match.style else None,
        }

    try:
        summary = _atomic_write(target_path, kind, _do_write)
    except Exception as exc:
        return _write_error(kind, f"write failed: {exc.__class__.__name__}: {exc}", resolved)
    return _write_result(resolved, target_path, kind, summary)
