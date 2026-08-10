"""office.* — read, outline, and edit tools for Word (.docx) and Excel (.xlsx).

Phase 1 (already shipped) — read-only:
  * office.read       — full content as LLM-ready markdown
  * office.outline    — cheap structural summary

Phase 2 (this module) — edits via COPY-ON-FIRST-WRITE:
  Docx: find_replace, insert_paragraph, update_paragraph,
        delete_paragraph, apply_style
  Xlsx: set_cells, set_formula, set_cell_format, merge_cells,
        add_row, read_range

The user's hard constraint for Phase 2: the original file is NEVER
modified. Every write tool copies the original to a sibling
`<basename>.edited.<ext>` on first write and targets the copy
thereafter. Repeated edits land in the same copy. The user can reset
edits by deleting `xxx.edited.docx` in Finder.

Because the original is untouched, rollback remains predictable. The Runtime
still reviews each concrete write operation before execution.

All write tools also use an atomic write pattern: write to
`<target>.tmp`, re-open with the appropriate library to verify it's a
valid OOXML file, then `os.replace(tmp, target)`. A mid-write failure
leaves `target` exactly as it was (the last known-good edit).

Path safety: read tools accept either a path under the authorized workspace
or an exact `/attachments/...` route owned by the current Run. Write tools
remain restricted to the authorized workspace and never mutate attachments.
"""

from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from docx.document import Document as DocxDocumentType
from docx.text.paragraph import Paragraph as DocxParagraph
from langchain_core.tools import tool
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import column_index_from_string, get_column_letter
from openpyxl.utils.cell import range_boundaries

from ..document_markdown import document_to_markdown
from .office_common import (
    MARKDOWN_CHAR_CAP as _MARKDOWN_CHAR_CAP,
)
from .office_common import (
    atomic_write as _atomic_write,
)
from .office_common import (
    ensure_copy_for_write as _ensure_copy_for_write,
)
from .office_common import (
    resolve_read_source as _resolve_read_source,
)
from .office_common import (
    resolve_write_target as _resolve_write_target,
)
from .office_common import (
    validate_path as _validate_path,
)
from .office_common import (
    write_error as _write_error,
)
from .office_common import (
    write_result as _write_result,
)
from .office_pptx import (
    _outline_pptx,
    office_add_image_to_slide,
    office_add_slide,
    office_create_pptx,
    office_delete_slide,
    office_read_slides,
    office_reorder_slides,
    office_set_slide_bullets,
    office_set_slide_notes,
    office_set_slide_title,
    office_update_slide,
)

# ═══════════════════════════════════════════════════════════════════════
# READ tools (Phase 1)
# ═══════════════════════════════════════════════════════════════════════


@tool("office.read")
def office_read(path: str) -> dict[str, Any]:
    """Read a Word (.docx), Excel (.xlsx), or PowerPoint (.pptx) file as markdown.

    Converts headings, paragraphs, tables, and cells into compact markdown
    the LLM can reason about directly.

    Does NOT open the right-side document preview panel. If you want the
    user to see the file rendered, mention the filename in your reply
    so the renderer makes it clickable. The preview opens when the user
    clicks, or after a successful office.* WRITE tool completes.

    Args:
        path: Exact Runtime `/attachments/...` route, or an existing absolute
              path under the authorized workspace.

    Returns:
        dict with keys:
          ok ("true" / "false")
          path (echoed back, absolute)
          kind ("word", "excel", or "powerpoint")
          markdown (the converted markdown content, possibly truncated)
          truncated ("true" / "false") — set when content exceeded the cap
          error (only present when ok="false")
    """
    source, kind, err = _resolve_read_source(path)
    if err is not None:
        return {"ok": "false", "error": err}
    assert source is not None and kind is not None  # for type checker
    reported_path = path if isinstance(source, bytes) else source
    try:
        text, truncated = document_to_markdown(
            source,
            Path(path).suffix.lower(),
            char_limit=_MARKDOWN_CHAR_CAP,
        )
    except Exception as exc:
        return {
            "ok": "false",
            "path": reported_path,
            "kind": kind,
            "error": f"failed to convert {kind}: {exc.__class__.__name__}: {exc}",
        }
    if truncated:
        text = text[:_MARKDOWN_CHAR_CAP] + "\n\n…(truncated)"
    return {
        "ok": "true",
        "path": reported_path,
        "kind": kind,
        "markdown": text,
        "truncated": "true" if truncated else "false",
    }


def _outline_docx(path: Any) -> dict[str, Any]:
    """Cheap structural summary of a .docx — heading text and paragraph count."""
    doc = DocxDocument(path)
    headings: list[dict[str, Any]] = []
    paragraph_count = 0
    table_count = len(doc.tables)
    for para in doc.paragraphs:
        paragraph_count += 1
        style_name = (para.style.name if para.style is not None else "") or ""
        if style_name.startswith("Heading"):
            level_str = style_name.split()[-1] if " " in style_name else "1"
            try:
                level = int(level_str)
            except ValueError:
                level = 1
            text = para.text.strip()
            if text:
                headings.append({"level": level, "text": text})
    return {
        "headings": headings,
        "paragraph_count": paragraph_count,
        "table_count": table_count,
    }


def _outline_xlsx(path: Any) -> dict[str, Any]:
    """Cheap structural summary of a .xlsx — sheet names + dimensions."""
    wb = load_workbook(path, read_only=True, data_only=True)
    sheets: list[dict[str, Any]] = []
    try:
        for name in wb.sheetnames:
            ws = wb[name]
            sheets.append(
                {
                    "name": name,
                    "rows": ws.max_row or 0,
                    "columns": ws.max_column or 0,
                }
            )
    finally:
        wb.close()
    return {"sheets": sheets}


@tool("office.outline")
def office_outline(path: str) -> dict[str, Any]:
    """Return a cheap structural summary of a .docx, .xlsx, or .pptx file.

    Use this BEFORE `office.read` when the file is large and you only need
    to know what's in it — for example: "tell me what sheets are in
    Q4.xlsx", "does report.docx have a section about pricing?", or "how
    many slides does pitch.pptx have?". Reading the outline is
    O(metadata); reading the full markdown is O(file).

    Args:
        path: Exact Runtime `/attachments/...` route, or an existing absolute
              path under the authorized workspace.

    Returns:
        dict with keys:
          ok ("true" / "false")
          path (echoed back, absolute)
          kind ("word", "excel", or "powerpoint")
          For .docx: headings, paragraph_count, table_count.
          For .xlsx: sheets (list of {name, rows, columns}).
          For .pptx: slides (list of {index, layout, title, bullets,
                     notes, shape_count, image_count}), slide_count.
          error (only present when ok="false")
    """
    source, kind, err = _resolve_read_source(path)
    if err is not None:
        return {"ok": "false", "error": err}
    assert source is not None and kind is not None
    reported_path = path if isinstance(source, bytes) else source
    parser_input = io.BytesIO(source) if isinstance(source, bytes) else source
    try:
        if kind == "word":
            details = _outline_docx(parser_input)
        elif kind == "excel":
            details = _outline_xlsx(parser_input)
        else:  # powerpoint
            details = _outline_pptx(parser_input)
    except Exception as exc:
        return {
            "ok": "false",
            "path": reported_path,
            "kind": kind,
            "error": f"failed to read {kind} outline: {exc.__class__.__name__}: {exc}",
        }
    return {
        "ok": "true",
        "path": reported_path,
        "kind": kind,
        **details,
    }


# ═══════════════════════════════════════════════════════════════════════
# DOCX write helpers
# ═══════════════════════════════════════════════════════════════════════


def _find_paragraph(doc: DocxDocumentType, target: str) -> DocxParagraph | None:
    """Return the first paragraph whose .text contains `target` (exact
    substring match), or None. Tables are walked too — paragraphs inside
    cells should be addressable just like top-level body paragraphs."""
    if not target:
        return None
    for para in doc.paragraphs:
        if target in para.text:
            return para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if target in para.text:
                        return para
    return None


def _set_paragraph_text(para: DocxParagraph, new_text: str) -> None:
    """Replace a paragraph's text with `new_text`, collapsing all
    existing runs into one. This LOSES run-level formatting inside the
    paragraph (bold-in-middle, font color, etc.) but preserves the
    paragraph-level style (Heading 1, Quote, list bullet, etc.). For
    the find_replace / update_paragraph use cases this is the right
    trade — agents are editing content, not micro-formatting."""
    if para.runs:
        # Reuse the first run to keep its char formatting as the new
        # paragraph's font (best-effort), then clear the rest.
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(new_text)


def _insert_paragraph_relative(
    anchor: DocxParagraph, content: str, position: str, style: str | None
) -> DocxParagraph:
    """Insert a new paragraph immediately before or after `anchor`.

    python-docx has no built-in "insert before/after" — we manipulate
    the XML directly. This is the standard recipe from the docx-py
    cookbook.
    """
    from copy import deepcopy

    new_p = deepcopy(anchor._p)
    # Strip the existing runs from the cloned XML so we get a blank
    # paragraph element with the same pPr (paragraph properties).
    for child in list(new_p):
        if child.tag.endswith("}r"):
            new_p.remove(child)
    if position == "before":
        anchor._p.addprevious(new_p)
    else:  # default + 'after'
        anchor._p.addnext(new_p)
    new_para = DocxParagraph(new_p, anchor._parent)
    if style:
        new_para.style = style
    new_para.add_run(content)
    return new_para


# ═══════════════════════════════════════════════════════════════════════
# DOCX write tools
# ═══════════════════════════════════════════════════════════════════════


@tool("office.find_replace")
def office_find_replace(
    path: str,
    find: str,
    replace: str,
    count: int | None = None,
) -> dict[str, Any]:
    """Replace text across a .docx. Writes to a copy named `<basename>.edited.docx`.

    The original file is NEVER modified — you always operate on a copy.
    If the copy already exists from a previous edit, this writes to
    that same copy (chained edits land in one file).

    Args:
        path: .docx to edit. May be the original or its `.edited` copy
              (idempotent: if you pass the edited copy, we write back
              to it).
        find: text to search for (exact substring match).
        replace: text to substitute in.
        count: optional — stop after this many replacements. None =
               replace all.

    Returns:
        dict with `ok`, `original_path`, `edited_path`, `kind`,
        `summary={replaced}`. Use `edited_path` for subsequent reads
        or edits on this document.

    Note on formatting: this tool collapses run-level char formatting
    in each modified paragraph (paragraph-level style is preserved).
    If you need to keep "bold middle word" intact, use
    `office.update_paragraph` instead — it's whole-paragraph anyway.
    """
    resolved, kind, err = _validate_path(path)
    if err is not None:
        return _write_error(None, err)
    assert resolved is not None and kind is not None
    if kind != "word":
        return _write_error(kind, "office.find_replace requires a .docx file", resolved)
    if not find:
        return _write_error(kind, "find text required", resolved)

    target = _ensure_copy_for_write(resolved)
    remaining = count if (count is None or count > 0) else 0

    def _do_write(tmp_path: str) -> dict[str, Any]:
        nonlocal remaining
        doc = DocxDocument(target)
        replaced = 0
        for para in _iter_all_paragraphs(doc):
            if remaining == 0:
                break
            text = para.text
            if find not in text:
                continue
            occurrences = text.count(find)
            if remaining is not None:
                occurrences = min(occurrences, remaining)
            new_text = text.replace(find, replace, occurrences)
            _set_paragraph_text(para, new_text)
            replaced += occurrences
            if remaining is not None:
                remaining -= occurrences
        doc.save(tmp_path)
        return {"replaced": replaced}

    try:
        summary = _atomic_write(target, kind, _do_write)
    except Exception as exc:
        return _write_error(kind, f"write failed: {exc.__class__.__name__}: {exc}", resolved)
    return _write_result(resolved, target, kind, summary)


def _iter_all_paragraphs(doc: DocxDocumentType):
    """Yield body paragraphs + paragraphs nested in table cells."""
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


@tool("office.insert_paragraph")
def office_insert_paragraph(
    path: str,
    anchor: str,
    content: str,
    position: str = "after",
    style: str | None = None,
) -> dict[str, Any]:
    """Insert a new paragraph relative to an anchor paragraph in a .docx.

    Writes to `<basename>.edited.docx`; original preserved.

    Args:
        path: .docx to edit.
        anchor: text that uniquely identifies the anchor paragraph
                (first paragraph whose text contains this substring).
                Empty string OR the literal "__END__" inserts at the
                very end of the document.
        content: text of the new paragraph.
        position: "before" or "after" the anchor (default "after").
                  Ignored when anchor is "__END__".
        style: optional python-docx style name (e.g. "Heading 1",
               "Normal", "Quote", "List Bullet"). Use `office.outline`
               to see which styles already appear in the document.

    Returns:
        dict with `ok`, `original_path`, `edited_path`, `kind`,
        `summary={inserted_at, style?}`.
    """
    resolved, kind, err = _validate_path(path)
    if err is not None:
        return _write_error(None, err)
    assert resolved is not None and kind is not None
    if kind != "word":
        return _write_error(kind, "office.insert_paragraph requires a .docx file", resolved)
    if position not in {"before", "after"}:
        return _write_error(
            kind, f"position must be 'before' or 'after', got {position!r}", resolved
        )

    target = _ensure_copy_for_write(resolved)

    def _do_write(tmp_path: str) -> dict[str, Any]:
        doc = DocxDocument(target)
        if anchor in ("", "__END__"):
            new_para = (
                doc.add_paragraph(content, style=style) if style else doc.add_paragraph(content)
            )
            doc.save(tmp_path)
            return {"inserted_at": "end", "style": new_para.style.name if new_para.style else None}
        match = _find_paragraph(doc, anchor)
        if match is None:
            raise ValueError(f"anchor not found: {anchor!r}")
        new_para = _insert_paragraph_relative(match, content, position, style)
        doc.save(tmp_path)
        return {
            "inserted_at": f"{position} anchor",
            "anchor_excerpt": match.text[:60],
            "style": new_para.style.name if new_para.style else None,
        }

    try:
        summary = _atomic_write(target, kind, _do_write)
    except Exception as exc:
        return _write_error(kind, f"write failed: {exc.__class__.__name__}: {exc}", resolved)
    return _write_result(resolved, target, kind, summary)


@tool("office.update_paragraph")
def office_update_paragraph(
    path: str,
    target: str,
    content: str,
    style: str | None = None,
) -> dict[str, Any]:
    """Rewrite the first paragraph that contains `target` substring.

    Writes to `<basename>.edited.docx`; original preserved.

    Args:
        path: .docx to edit.
        target: substring that identifies the paragraph to update
                (first match wins; pick something unique).
        content: new full text of that paragraph (replaces ALL existing
                 text in the paragraph).
        style: optional new paragraph style name.

    Returns:
        dict with `ok`, `original_path`, `edited_path`, `kind`,
        `summary={updated_excerpt, style?}`.
    """
    resolved, kind, err = _validate_path(path)
    if err is not None:
        return _write_error(None, err)
    assert resolved is not None and kind is not None
    if kind != "word":
        return _write_error(kind, "office.update_paragraph requires a .docx file", resolved)
    if not target:
        return _write_error(kind, "target text required", resolved)

    target_path = _ensure_copy_for_write(resolved)

    def _do_write(tmp_path: str) -> dict[str, Any]:
        doc = DocxDocument(target_path)
        match = _find_paragraph(doc, target)
        if match is None:
            raise ValueError(f"target paragraph not found: {target!r}")
        old_excerpt = match.text[:60]
        _set_paragraph_text(match, content)
        if style:
            match.style = style
        doc.save(tmp_path)
        return {
            "updated_excerpt": old_excerpt,
            "style": match.style.name if match.style else None,
        }

    try:
        summary = _atomic_write(target_path, kind, _do_write)
    except Exception as exc:
        return _write_error(kind, f"write failed: {exc.__class__.__name__}: {exc}", resolved)
    return _write_result(resolved, target_path, kind, summary)


@tool("office.delete_paragraph")
def office_delete_paragraph(path: str, target: str) -> dict[str, Any]:
    """Delete the first paragraph that contains `target` substring.

    Writes to `<basename>.edited.docx`; original preserved.

    Args:
        path: .docx to edit.
        target: substring that identifies the paragraph (first match wins).

    Returns:
        dict with `ok`, `original_path`, `edited_path`, `kind`,
        `summary={deleted_excerpt}`.
    """
    resolved, kind, err = _validate_path(path)
    if err is not None:
        return _write_error(None, err)
    assert resolved is not None and kind is not None
    if kind != "word":
        return _write_error(kind, "office.delete_paragraph requires a .docx file", resolved)
    if not target:
        return _write_error(kind, "target text required", resolved)

    target_path = _ensure_copy_for_write(resolved)

    def _do_write(tmp_path: str) -> dict[str, Any]:
        doc = DocxDocument(target_path)
        match = _find_paragraph(doc, target)
        if match is None:
            raise ValueError(f"target paragraph not found: {target!r}")
        deleted_excerpt = match.text[:60]
        # Detach the paragraph element from its parent (python-docx has
        # no public `.delete()`; we manipulate XML directly).
        p_element = match._p
        p_element.getparent().remove(p_element)
        doc.save(tmp_path)
        return {"deleted_excerpt": deleted_excerpt}

    try:
        summary = _atomic_write(target_path, kind, _do_write)
    except Exception as exc:
        return _write_error(kind, f"write failed: {exc.__class__.__name__}: {exc}", resolved)
    return _write_result(resolved, target_path, kind, summary)


@tool("office.apply_style")
def office_apply_style(path: str, target: str, style: str) -> dict[str, Any]:
    """Change the paragraph style of the first paragraph matching `target`.

    Writes to `<basename>.edited.docx`; original preserved.

    Args:
        path: .docx to edit.
        target: substring identifying the paragraph (first match wins).
        style: python-docx style name. Common values: "Heading 1",
               "Heading 2", "Heading 3", "Normal", "Quote", "Title",
               "Subtitle", "List Bullet", "List Number".

    Returns:
        dict with `ok`, `original_path`, `edited_path`, `kind`,
        `summary={paragraph_excerpt, old_style, new_style}`.
    """
    resolved, kind, err = _validate_path(path)
    if err is not None:
        return _write_error(None, err)
    assert resolved is not None and kind is not None
    if kind != "word":
        return _write_error(kind, "office.apply_style requires a .docx file", resolved)
    if not target or not style:
        return _write_error(kind, "target and style are both required", resolved)

    target_path = _ensure_copy_for_write(resolved)

    def _do_write(tmp_path: str) -> dict[str, Any]:
        doc = DocxDocument(target_path)
        match = _find_paragraph(doc, target)
        if match is None:
            raise ValueError(f"target paragraph not found: {target!r}")
        old_style = match.style.name if match.style else None
        match.style = style
        doc.save(tmp_path)
        return {
            "paragraph_excerpt": match.text[:60],
            "old_style": old_style,
            "new_style": match.style.name if match.style else None,
        }

    try:
        summary = _atomic_write(target_path, kind, _do_write)
    except Exception as exc:
        return _write_error(kind, f"write failed: {exc.__class__.__name__}: {exc}", resolved)
    return _write_result(resolved, target_path, kind, summary)


# ═══════════════════════════════════════════════════════════════════════
# XLSX write helpers
# ═══════════════════════════════════════════════════════════════════════


_CELL_REF_RE = re.compile(r"^([A-Za-z]+)(\d+)$")


def _parse_cell_ref(ref: str) -> tuple[int, int]:
    """Parse an A1-style cell ref into (row, col) 1-indexed tuple."""
    m = _CELL_REF_RE.match(ref.strip())
    if not m:
        raise ValueError(f"invalid cell reference: {ref!r} (expected like 'A1')")
    col_str, row_str = m.groups()
    return int(row_str), column_index_from_string(col_str.upper())


def _parse_range(range_ref: str) -> tuple[int, int, int, int]:
    """Parse a range like 'A1:C3' (or 'A1' for a single cell) into
    (min_row, min_col, max_row, max_col), all 1-indexed."""
    ref = range_ref.strip()
    if ":" not in ref:
        # Single cell — treat as a 1x1 range.
        row, col = _parse_cell_ref(ref)
        return row, col, row, col
    min_col, min_row, max_col, max_row = range_boundaries(ref.upper())
    return min_row, min_col, max_row, max_col


def _resolve_sheet(wb, sheet: str | None):
    """Return the worksheet by name, or the active sheet when sheet is None."""
    if not sheet:
        return wb.active
    if sheet not in wb.sheetnames:
        raise ValueError(f"sheet not found: {sheet!r}; available: {wb.sheetnames}")
    return wb[sheet]


def _color_to_hex(color: str) -> str:
    """Normalize a CSS-ish color string into openpyxl's 8-char AARRGGBB
    or 6-char RRGGBB hex (without leading '#'). Accepts '#RRGGBB',
    'RRGGBB', '#AARRGGBB', 'AARRGGBB'. Bare names not supported."""
    c = color.strip().lstrip("#").upper()
    if len(c) not in (6, 8) or not all(ch in "0123456789ABCDEF" for ch in c):
        raise ValueError(f"color must be hex like '#FF5722' or '#80FF5722', got {color!r}")
    return c


# ═══════════════════════════════════════════════════════════════════════
# XLSX write tools
# ═══════════════════════════════════════════════════════════════════════


@tool("office.set_cells")
def office_set_cells(
    path: str,
    sheet: str | None,
    range: str,
    values: list[list[Any]],
) -> dict[str, Any]:
    """Write a 2D values block into a rectangular cell range.

    Writes to `<basename>.edited.xlsx`; original preserved.

    Args:
        path: .xlsx to edit.
        sheet: sheet name (or None for the active sheet).
        range: A1-style range like "A1:C3" or single cell "B5".
        values: 2D list shaped to match the range. Top-left is range
                top-left. Each inner list is one row. Extra cells in
                the range that lack values stay unchanged; extra
                values beyond the range are ignored.

    Returns:
        dict with `ok`, `original_path`, `edited_path`, `kind`,
        `summary={range, rows_written, cols_written, cells_written}`.
    """
    resolved, kind, err = _validate_path(path)
    if err is not None:
        return _write_error(None, err)
    assert resolved is not None and kind is not None
    if kind != "excel":
        return _write_error(kind, "office.set_cells requires a .xlsx file", resolved)
    if not values or not isinstance(values, list):
        return _write_error(kind, "values must be a non-empty 2D list", resolved)

    target_path = _ensure_copy_for_write(resolved)

    def _do_write(tmp_path: str) -> dict[str, Any]:
        wb = load_workbook(target_path)
        try:
            ws = _resolve_sheet(wb, sheet)
            min_row, min_col, max_row, max_col = _parse_range(range)
            rng_rows = max_row - min_row + 1
            rng_cols = max_col - min_col + 1
            written = 0
            for r_offset, row_values in enumerate(values):
                if r_offset >= rng_rows:
                    break
                if not isinstance(row_values, list):
                    raise ValueError(f"values[{r_offset}] must be a list")
                for c_offset, cell_value in enumerate(row_values):
                    if c_offset >= rng_cols:
                        break
                    ws.cell(row=min_row + r_offset, column=min_col + c_offset, value=cell_value)
                    written += 1
            wb.save(tmp_path)
        finally:
            wb.close()
        return {
            "range": range,
            "rows_written": min(len(values), rng_rows),
            "cols_written": rng_cols,
            "cells_written": written,
        }

    try:
        summary = _atomic_write(target_path, kind, _do_write)
    except Exception as exc:
        return _write_error(kind, f"write failed: {exc.__class__.__name__}: {exc}", resolved)
    return _write_result(resolved, target_path, kind, summary)


@tool("office.set_formula")
def office_set_formula(
    path: str,
    sheet: str | None,
    cell: str,
    formula: str,
) -> dict[str, Any]:
    """Write a formula into one cell. Writes to `<basename>.edited.xlsx`.

    NOTE: openpyxl writes the formula text — it does NOT evaluate it.
    The cell's displayed value updates only when Microsoft Excel /
    LibreOffice / Numbers opens the file. The right-side preview shows
    the literal formula text until then. This is an openpyxl limitation,
    not a bug in this tool.

    Args:
        path: .xlsx to edit.
        sheet: sheet name (None = active).
        cell: target cell like "D2".
        formula: must start with "=", e.g. "=SUM(A2:A10)",
                 "=IF(B2>0, \\"Pos\\", \\"Neg\\")", "=VLOOKUP(...)".

    Returns:
        dict with `ok`, `original_path`, `edited_path`, `kind`,
        `summary={cell, formula}`.
    """
    resolved, kind, err = _validate_path(path)
    if err is not None:
        return _write_error(None, err)
    assert resolved is not None and kind is not None
    if kind != "excel":
        return _write_error(kind, "office.set_formula requires a .xlsx file", resolved)
    formula = (formula or "").strip()
    if not formula.startswith("="):
        return _write_error(kind, "formula must start with '='", resolved)

    target_path = _ensure_copy_for_write(resolved)

    def _do_write(tmp_path: str) -> dict[str, Any]:
        wb = load_workbook(target_path)
        try:
            ws = _resolve_sheet(wb, sheet)
            row, col = _parse_cell_ref(cell)
            ws.cell(row=row, column=col, value=formula)
            wb.save(tmp_path)
        finally:
            wb.close()
        return {"cell": cell, "formula": formula}

    try:
        summary = _atomic_write(target_path, kind, _do_write)
    except Exception as exc:
        return _write_error(kind, f"write failed: {exc.__class__.__name__}: {exc}", resolved)
    return _write_result(resolved, target_path, kind, summary)


@tool("office.set_cell_format")
def office_set_cell_format(
    path: str,
    sheet: str | None,
    range: str,
    bold: bool | None = None,
    italic: bool | None = None,
    font_size: float | None = None,
    font_color: str | None = None,
    bg_color: str | None = None,
    align: str | None = None,
    border: bool | None = None,
) -> dict[str, Any]:
    """Apply font / fill / alignment / border formatting to a cell range.

    Writes to `<basename>.edited.xlsx`; original preserved. Only the
    arguments you pass are applied — None means "leave unchanged".

    Args:
        path: .xlsx to edit.
        sheet: sheet name (None = active).
        range: A1-style range.
        bold, italic: True / False (None to leave alone).
        font_size: e.g. 14.
        font_color, bg_color: hex like "#FF5722" or "FF5722". Alpha
                              prefix is accepted ("#80FF5722").
        align: one of "left" / "center" / "right" (horizontal only).
        border: True adds a thin black border to every cell in the
                range; False clears borders.

    Returns:
        dict with `ok`, `original_path`, `edited_path`, `kind`,
        `summary={range, cells_formatted, applied}` where `applied`
        is the list of attribute names that were actually changed.
    """
    resolved, kind, err = _validate_path(path)
    if err is not None:
        return _write_error(None, err)
    assert resolved is not None and kind is not None
    if kind != "excel":
        return _write_error(kind, "office.set_cell_format requires a .xlsx file", resolved)

    applied: list[str] = []
    try:
        font_color_hex = _color_to_hex(font_color) if font_color else None
        bg_color_hex = _color_to_hex(bg_color) if bg_color else None
    except ValueError as exc:
        return _write_error(kind, str(exc), resolved)
    if align is not None and align not in {"left", "center", "right"}:
        return _write_error(kind, f"align must be left/center/right, got {align!r}", resolved)

    target_path = _ensure_copy_for_write(resolved)

    def _do_write(tmp_path: str) -> dict[str, Any]:
        wb = load_workbook(target_path)
        try:
            ws = _resolve_sheet(wb, sheet)
            min_row, min_col, max_row, max_col = _parse_range(range)
            count = 0
            side = Side(border_style="thin", color="FF000000") if border else None
            border_style = (
                Border(left=side, right=side, top=side, bottom=side)
                if border
                else (Border() if border is False else None)
            )
            for r in range_iter(min_row, max_row):
                for c in range_iter(min_col, max_col):
                    cell_obj = ws.cell(row=r, column=c)
                    # Font — clone current font then patch fields the
                    # caller asked to change.
                    if any(v is not None for v in (bold, italic, font_size, font_color_hex)):
                        existing = cell_obj.font
                        cell_obj.font = Font(
                            name=existing.name,
                            size=font_size if font_size is not None else existing.size,
                            bold=bold if bold is not None else existing.bold,
                            italic=italic if italic is not None else existing.italic,
                            color=font_color_hex if font_color_hex is not None else existing.color,
                        )
                        for f in ("bold", "italic", "font_size", "font_color"):
                            if locals().get(f) is not None or (
                                f == "font_color" and font_color_hex
                            ):
                                if f not in applied:
                                    applied.append(f)
                    if bg_color_hex is not None:
                        cell_obj.fill = PatternFill(
                            fill_type="solid", start_color=bg_color_hex, end_color=bg_color_hex
                        )
                        if "bg_color" not in applied:
                            applied.append("bg_color")
                    if align is not None:
                        existing = cell_obj.alignment
                        cell_obj.alignment = Alignment(
                            horizontal=align,
                            vertical=existing.vertical,
                            wrap_text=existing.wrap_text,
                        )
                        if "align" not in applied:
                            applied.append("align")
                    if border_style is not None:
                        cell_obj.border = border_style
                        if "border" not in applied:
                            applied.append("border")
                    count += 1
            wb.save(tmp_path)
        finally:
            wb.close()
        return {"range": range, "cells_formatted": count, "applied": applied}

    try:
        summary = _atomic_write(target_path, kind, _do_write)
    except Exception as exc:
        return _write_error(kind, f"write failed: {exc.__class__.__name__}: {exc}", resolved)
    return _write_result(resolved, target_path, kind, summary)


def range_iter(lo: int, hi: int):
    """Inclusive integer range — `range(lo, hi+1)` shorthand."""
    return range(lo, hi + 1)


@tool("office.merge_cells")
def office_merge_cells(
    path: str,
    sheet: str | None,
    range: str,
) -> dict[str, Any]:
    """Merge a rectangular range into a single cell.

    Writes to `<basename>.edited.xlsx`; original preserved.

    Args:
        path: .xlsx to edit.
        sheet: sheet name (None = active).
        range: A1-style range like "A1:C1" (header row span) or
               "A1:B4" (vertical+horizontal merge).

    Returns:
        dict with `ok`, `original_path`, `edited_path`, `kind`,
        `summary={range}`.
    """
    resolved, kind, err = _validate_path(path)
    if err is not None:
        return _write_error(None, err)
    assert resolved is not None and kind is not None
    if kind != "excel":
        return _write_error(kind, "office.merge_cells requires a .xlsx file", resolved)
    if ":" not in range:
        return _write_error(
            kind, "merge_cells requires a range like 'A1:C1', not a single cell", resolved
        )

    target_path = _ensure_copy_for_write(resolved)

    def _do_write(tmp_path: str) -> dict[str, Any]:
        wb = load_workbook(target_path)
        try:
            ws = _resolve_sheet(wb, sheet)
            ws.merge_cells(range)
            wb.save(tmp_path)
        finally:
            wb.close()
        return {"range": range}

    try:
        summary = _atomic_write(target_path, kind, _do_write)
    except Exception as exc:
        return _write_error(kind, f"write failed: {exc.__class__.__name__}: {exc}", resolved)
    return _write_result(resolved, target_path, kind, summary)


@tool("office.add_row")
def office_add_row(
    path: str,
    sheet: str | None,
    values: list[Any],
    position: str | int = "append",
) -> dict[str, Any]:
    """Insert a row into a sheet.

    Writes to `<basename>.edited.xlsx`; original preserved.

    Args:
        path: .xlsx to edit.
        sheet: sheet name (None = active).
        values: list of cell values; the row is filled left-to-right
                from column A.
        position: "append" (default — adds after the last used row)
                  or an integer row number (1-indexed) to insert BEFORE
                  that row, shifting subsequent rows down.

    Returns:
        dict with `ok`, `original_path`, `edited_path`, `kind`,
        `summary={row, cells_written}`.
    """
    resolved, kind, err = _validate_path(path)
    if err is not None:
        return _write_error(None, err)
    assert resolved is not None and kind is not None
    if kind != "excel":
        return _write_error(kind, "office.add_row requires a .xlsx file", resolved)
    if not isinstance(values, list):
        return _write_error(kind, "values must be a list", resolved)
    if position != "append" and not isinstance(position, int):
        return _write_error(kind, "position must be 'append' or an integer row number", resolved)

    target_path = _ensure_copy_for_write(resolved)

    def _do_write(tmp_path: str) -> dict[str, Any]:
        wb = load_workbook(target_path)
        try:
            ws = _resolve_sheet(wb, sheet)
            if position == "append":
                ws.append(values)
                row_idx = ws.max_row
            else:
                row_idx = int(position)
                if row_idx < 1:
                    raise ValueError(f"row position must be >= 1, got {row_idx}")
                ws.insert_rows(row_idx)
                for i, v in enumerate(values, start=1):
                    ws.cell(row=row_idx, column=i, value=v)
            wb.save(tmp_path)
        finally:
            wb.close()
        return {"row": row_idx, "cells_written": len(values)}

    try:
        summary = _atomic_write(target_path, kind, _do_write)
    except Exception as exc:
        return _write_error(kind, f"write failed: {exc.__class__.__name__}: {exc}", resolved)
    return _write_result(resolved, target_path, kind, summary)


@tool("office.read_range")
def office_read_range(
    path: str,
    sheet: str | None,
    range: str,
) -> dict[str, Any]:
    """Read a specific cell range from a .xlsx as structured JSON.

    This is the precise complement to `office.read` (which dumps the
    whole workbook as markdown). Use this when you only need a subset
    — much smaller LLM context spend, and you get raw types (int /
    float / str / bool / None) plus the formula text for formula
    cells.

    Args:
        path: Runtime `/attachments/...` route or workspace .xlsx path.
        sheet: sheet name (None = active).
        range: A1-style range like "A1:C10" or single cell "B5".

    Returns:
        dict with:
          ok
          path, kind
          sheet (resolved sheet name)
          range (echoed back)
          values   — 2D list, computed values (formulas → cached result)
          formulas — 2D list, formula text where present, else None
    """
    source, kind, err = _resolve_read_source(path)
    if err is not None:
        return {"ok": "false", "error": err}
    assert source is not None and kind is not None
    reported_path = path if isinstance(source, bytes) else source
    if kind != "excel":
        return {
            "ok": "false",
            "error": "office.read_range requires a .xlsx file",
            "path": reported_path,
        }
    try:
        # Two passes: data_only=True for cached values, data_only=False
        # for formula text. openpyxl can't give both in one open.
        wb_values = load_workbook(
            io.BytesIO(source) if isinstance(source, bytes) else source,
            read_only=True,
            data_only=True,
        )
        wb_formulas = load_workbook(
            io.BytesIO(source) if isinstance(source, bytes) else source,
            read_only=True,
            data_only=False,
        )
    except Exception as exc:
        return {
            "ok": "false",
            "path": reported_path,
            "kind": kind,
            "error": f"failed to open .xlsx: {exc.__class__.__name__}: {exc}",
        }
    try:
        ws_values = _resolve_sheet(wb_values, sheet)
        ws_formulas = _resolve_sheet(wb_formulas, sheet)
        min_row, min_col, max_row, max_col = _parse_range(range)
        values_grid: list[list[Any]] = []
        formulas_grid: list[list[str | None]] = []
        for r in range_iter(min_row, max_row):
            v_row: list[Any] = []
            f_row: list[str | None] = []
            for c in range_iter(min_col, max_col):
                v_row.append(ws_values.cell(row=r, column=c).value)
                f_cell = ws_formulas.cell(row=r, column=c).value
                f_row.append(f_cell if isinstance(f_cell, str) and f_cell.startswith("=") else None)
            values_grid.append(v_row)
            formulas_grid.append(f_row)
    except Exception as exc:
        wb_values.close()
        wb_formulas.close()
        return {
            "ok": "false",
            "path": reported_path,
            "kind": kind,
            "error": f"failed to read range {range!r}: {exc.__class__.__name__}: {exc}",
        }
    sheet_name = ws_values.title
    wb_values.close()
    wb_formulas.close()
    return {
        "ok": "true",
        "path": reported_path,
        "kind": kind,
        "sheet": sheet_name,
        "range": range,
        "values": values_grid,
        "formulas": formulas_grid,
    }


# Phase 1 surface (read-only).
OFFICE_READ_TOOLS = [office_read, office_outline, office_read_range, office_read_slides]

# Phase 2 + Phase 3 surface (writes, all copy-on-first-write — except
# office.create_pptx which produces the original file).
OFFICE_WRITE_TOOLS = [
    office_find_replace,
    office_insert_paragraph,
    office_update_paragraph,
    office_delete_paragraph,
    office_apply_style,
    office_set_cells,
    office_set_formula,
    office_set_cell_format,
    office_merge_cells,
    office_add_row,
    office_create_pptx,
    office_add_slide,
    office_update_slide,
    office_delete_slide,
    office_reorder_slides,
    office_set_slide_title,
    office_set_slide_bullets,
    office_set_slide_notes,
    office_add_image_to_slide,
]


# Re-exported for tests that want to construct paths the same way the
# tools do without poking at the private helper.
def edited_copy_path(original_path: str) -> str:
    return _resolve_write_target(original_path)


# Silence unused-import warnings for symbols re-exported by tests.
__all__ = [
    "OFFICE_READ_TOOLS",
    "OFFICE_WRITE_TOOLS",
    "edited_copy_path",
    "get_column_letter",  # tests construct range refs
    "office_add_image_to_slide",
    "office_add_row",
    "office_add_slide",
    "office_apply_style",
    "office_create_pptx",
    "office_delete_paragraph",
    "office_delete_slide",
    "office_find_replace",
    "office_insert_paragraph",
    "office_merge_cells",
    "office_outline",
    "office_read",
    "office_read_range",
    "office_read_slides",
    "office_reorder_slides",
    "office_set_cell_format",
    "office_set_cells",
    "office_set_formula",
    "office_set_slide_bullets",
    "office_set_slide_notes",
    "office_set_slide_title",
    "office_update_paragraph",
    "office_update_slide",
]
